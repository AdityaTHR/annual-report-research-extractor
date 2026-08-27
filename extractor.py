import csv
import html
import io
import json
import re
import zipfile
from collections import Counter
from functools import lru_cache
from copy import deepcopy
from difflib import SequenceMatcher

try:
    import pymupdf as fitz
except Exception:
    import fitz
import pytesseract
from PIL import Image
from docx import Document

try:
    import pymupdf4llm
    HAVE_PYMUPDF4LLM = True
except Exception:
    pymupdf4llm = None
    HAVE_PYMUPDF4LLM = False

PRESETS = {
    "Chairman Message": {
        "aliases": [
            "message from the chairman", "chairman's message", "chairperson's message",
            "message from the chairperson", "letter from the chairman", "chairman's letter",
            "chairman's statement", "chairman statement", "statement from the chairman",
            "executive chairman's message", "message from the executive chairman",
        ],
    },
    "CEO Message": {
        "aliases": [
            "message from the ceo", "ceo message", "ceo's message", "letter from the ceo",
            "chief executive officer's message", "message from the chief executive officer",
            "letter from the chief executive officer",
        ],
    },
    "Managing Director Message": {
        "aliases": [
            "message from the managing director", "managing director's message",
            "md message", "letter from the managing director",
        ],
    },
    "Management Discussion & Analysis": {
        "aliases": [
            "management discussion & analysis report",
            "management discussion and analysis report",
            "management discussion & analysis",
            "management discussion and analysis",
            "management discussion analysis", "md&a",
        ],
    },
    "Business Responsibility & Sustainability Report (BRSR)": {
        "aliases": [
            "business responsibility & sustainability report",
            "business responsibility and sustainability report", "brsr report", "brsr",
        ],
    },
    "Business Responsibility Report (BRR)": {
        "aliases": ["business responsibility report", "business responsibility (br) report", "brr report"],
    },
    "ESG Report": {
        "aliases": ["esg report", "environmental social and governance report", "environment social and governance report"],
    },
    "Sustainability Report": {
        "aliases": ["sustainability report"],
    },
}

FORMAT_EXT = {"TXT": "txt", "JSON": "json", "PDF": "pdf", "DOCX": "docx", "CSV": "csv", "MD": "md"}

# V13 Generic Combined R3: structure-first, no company/year/page-specific rules.

@lru_cache(maxsize=32768)
def _compact(s):
    return re.sub(r"[^a-z0-9]+", "", s.lower())

def _norm_line(s):
    return re.sub(r"\s+", " ", s).strip()

def _ocr(page):
    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), colorspace=fitz.csGRAY, alpha=False)
    img = Image.frombytes("L", (pix.width, pix.height), pix.samples)
    return pytesseract.image_to_string(img, config="--psm 6")

def _layout_lines(page):
    out = []
    try:
        d = page.get_text("dict")
        order = 0
        for block in d.get("blocks", []):
            for line in block.get("lines", []):
                spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                x0 = min(s["bbox"][0] for s in spans)
                y0 = min(s["bbox"][1] for s in spans)
                x1 = max(s["bbox"][2] for s in spans)
                y1 = max(s["bbox"][3] for s in spans)
                fonts = [str(s.get("font", "")) for s in spans]
                flags = [int(s.get("flags", 0) or 0) for s in spans]
                out.append({
                    "order": order,
                    "text": text,
                    "bbox": [float(x0), float(y0), float(x1), float(y1)],
                    "size": float(max(s.get("size", 0) for s in spans)),
                    "bold": any("bold" in f.lower() for f in fonts) or any(f & 16 for f in flags),
                    "font": max(fonts, key=len) if fonts else "",
                })
                order += 1
    except Exception:
        return []
    return out

def _md_plain(text):
    if not text:
        return ""
    out = []
    for raw in str(text).splitlines():
        line = raw.strip()
        if not line:
            out.append("")
            continue
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"^[-*+]\s+", "", line)
        line = re.sub(r"^>\s?", "", line)
        line = line.replace("**", "").replace("__", "")
        line = re.sub(r"(?<!\w)[*_](.*?)[*_](?!\w)", r"\1", line)
        line = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", line)
        out.append(_norm_line(line))
    return re.sub(r"\n{3,}", "\n\n", "\n".join(out)).strip()

def _chunk_boxes(chunk):
    boxes = chunk.get("page_boxes")
    if boxes is None:
        boxes = (chunk.get("metadata") or {}).get("page_boxes")
    return boxes or []

def _layout_chunk_lines(chunk):
    md = chunk.get("text") or ""
    rows = []
    order = 0
    for box in _chunk_boxes(chunk):
        if not isinstance(box, dict):
            continue
        cls = str(box.get("class") or box.get("boxclass") or "text")
        if cls in {"page-header", "page-footer", "picture"}:
            continue
        pos = box.get("pos")
        text = ""
        if isinstance(pos, (list, tuple)) and len(pos) == 2:
            try:
                text = md[int(pos[0]):int(pos[1])]
            except Exception:
                text = ""
        if not text:
            text = box.get("text") or ""
        bbox = box.get("bbox")
        if bbox is not None:
            try:
                bbox = [float(x) for x in bbox]
            except Exception:
                bbox = None
        for piece in _md_plain(text).splitlines():
            piece = _norm_line(piece)
            if not piece:
                continue
            rows.append({
                "order": order,
                "text": piece,
                "bbox": bbox,
                "size": None,
                "bold": cls in {"title", "section-header"},
                "layout_class": cls,
            })
            order += 1
    return rows

def _text_quality(text):
    if not text:
        return 0.0
    s = str(text)
    printable = sum(ch.isprintable() or ch in "\n\t" for ch in s) / max(1, len(s))
    letters = sum(ch.isalpha() for ch in s)
    bad = sum(ch == "\ufffd" or (0xE000 <= ord(ch) <= 0xF8FF) for ch in s)
    odd = len(re.findall(r"[ƀ-ɏǀ-ǿ]{2,}|[\x00-\x08\x0b\x0c\x0e-\x1f]", s))
    if letters == 0:
        return 0.35 * printable
    corruption = min(1.0, (bad * 4 + odd * 6) / max(1, letters))
    return max(0.0, min(1.0, printable * (1.0 - corruption)))

def _layout_chunks(doc):
    if not HAVE_PYMUPDF4LLM:
        return None
    try:
        pymupdf4llm.use_layout(True)
        chunks = pymupdf4llm.to_markdown(
            doc, page_chunks=True, header=False, footer=False, ignore_images=True,
            write_images=False, embed_images=False, use_ocr=False, force_ocr=False, show_progress=False,
        )
        return chunks if isinstance(chunks, list) else None
    except Exception:
        return None

def extract_source(name, data):
    ext = name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        chunks = _layout_chunks(doc)

        structural_tokens = (
            "content", "chairman", "chief executive", "ceo", "managing director",
            "management discussion", "corporate governance", "responsibility",
            "sustainability", "directors' report", "directors’ report", "board's report",
            "boards’ report", "annexure", "financial statements", "clinical governance",
            "corporate social responsibility", "notice", "awards", "company overview",
            "about the company", "company profile", "at a glance", "performance highlights",
            "corporate review", "risk management", "financial highlights", "secretarial audit",
            "business overview", "human capital", "cybersecurity",
        )

        for i, page in enumerate(doc):
            native = page.get_text("text")
            text = native
            method = "native"
            native_lines = _layout_lines(page)
            for row in native_lines:
                row["layout_class"] = "native"
            lines = []

            chunk = chunks[i] if chunks and i < len(chunks) else None
            if chunk:
                layout_text = _md_plain(chunk.get("text") or "")
                layout_lines = _layout_chunk_lines(chunk)
                if len(layout_text.strip()) >= 20 and _text_quality(layout_text) >= 0.55:
                    text = layout_text
                    method = "layout"
                if layout_lines:
                    lines = layout_lines

            if not lines:
                low = native.lower()
                need_layout = i < 30 or any(tok in low for tok in structural_tokens)
                lines = native_lines if need_layout else [
                    {"order": j, "text": t, "bbox": None, "size": None, "layout_class": "text"}
                    for j, t in enumerate(text.splitlines()) if t.strip()
                ]

            native_edge = native_lines
            if native_edge:
                h = float(page.rect.height)
                edge_rows = []
                for row in native_edge:
                    bb = row.get("bbox")
                    if not bb:
                        continue
                    if bb[1] < 0.12 * h or bb[3] > 0.90 * h:
                        x = dict(row)
                        x["order"] = len(lines) + len(edge_rows)
                        x["layout_class"] = "folio-edge"
                        edge_rows.append(x)
                lines = lines + edge_rows

            if len(text.strip()) < 40 or _text_quality(text) < 0.45:
                ocr_text = _ocr(page)
                if len(ocr_text.strip()) > len(text.strip()) or _text_quality(ocr_text) > _text_quality(text):
                    text = ocr_text
                    method = "ocr"
                    lines = [
                        {"order": j, "text": t, "bbox": None, "size": None, "layout_class": "text"}
                        for j, t in enumerate(text.splitlines()) if t.strip()
                    ] + [r for r in lines if r.get("layout_class") == "folio-edge"]

            ocr_heading_lines = []
            low_native = native.lower().replace("’", "'")
            landscape_two_up = float(page.rect.width) / max(float(page.rect.height), 1.0) >= 1.25
            probe_role_heading = (
                i < 15 and landscape_two_up
                and ("dear shareholder" in low_native or "dear stakeholder" in low_native)
                and not any(k in low_native for k in ["chairman's message", "chairman's letter", "letter from the chairman", "letter from the ceo", "ceo message"])
            )
            if probe_role_heading:
                probe = _ocr(page)
                ocr_heading_lines = [
                    {"order": j, "text": t, "bbox": None, "size": None, "layout_class": "ocr-heading"}
                    for j, t in enumerate(probe.splitlines()) if _norm_line(t)
                ]

            pages.append({
                "page": i + 1, "text": text, "native_text": native, "method": method,
                "width": float(page.rect.width), "height": float(page.rect.height),
                "lines": lines, "native_lines": native_lines, "ocr_heading_lines": ocr_heading_lines,
                "layout_engine": bool(chunk),
            })
        return pages

    if ext in ("txt", "md"):
        text = data.decode("utf-8", errors="ignore")
        parts = re.split(r"(?m)^\s*=====\s*PAGE\s+(\d+)\s*=====\s*$", text)
        if len(parts) > 2:
            result = []
            for i in range(1, len(parts), 2):
                body = parts[i + 1]
                lines = [{"order": j, "text": t, "bbox": None, "size": None} for j, t in enumerate(body.splitlines()) if t.strip()]
                result.append({"page": int(parts[i]), "text": body, "method": "text", "lines": lines})
            return result
        lines = [{"order": j, "text": t, "bbox": None, "size": None} for j, t in enumerate(text.splitlines()) if t.strip()]
        return [{"page": None, "text": text, "method": "text", "lines": lines}]

    if ext == "docx":
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        lines = [{"order": j, "text": t, "bbox": None, "size": None} for j, t in enumerate(text.splitlines()) if t.strip()]
        return [{"page": None, "text": text, "method": "docx", "lines": lines}]

    raise ValueError(f"Unsupported file type: {ext}")


def clean_pages(pages):
    out = deepcopy(pages)
    if len(out) <= 1:
        out[0]["text"] = "\n".join(_norm_line(x) for x in out[0]["text"].splitlines() if _norm_line(x))
        return out

    edge = Counter()
    n = len(out)
    for p in out:
        lines = [_norm_line(x) for x in p["text"].splitlines() if _norm_line(x)]
        edge_lines = lines[:7] + lines[-7:]
        for x in set(edge_lines):
            words = re.findall(r"[A-Za-z]+", x)
            if 4 < len(x) <= 140 and len(words) >= 2:
                edge[x] += 1

    repeated = {x for x, c in edge.items() if c >= max(3, int(0.20 * n))}
    for p in out:
        original = [_norm_line(x) for x in p["text"].splitlines() if _norm_line(x)]
        cleaned = []
        last = len(original) - 1
        for i, x in enumerate(original):
            at_edge = i < 7 or i > last - 7
            if at_edge and x in repeated:
                continue
            cleaned.append(x)
        p["text"] = "\n".join(cleaned)
    return out


COMMON_SECTION_HEADINGS = [
    "chairman's message", "chairman's letter", "letter from the chairman",
    "letter from the ceo", "managing director's message",
    "directors' report", "board's report",
    "management discussion and analysis", "management discussion & analysis",
    "corporate governance report", "corporate social responsibility",
    "business responsibility report", "business responsibility and sustainability report",
    "standalone financial statements", "consolidated financial statements",
    "independent auditor's report", "clinical governance", "notice", "gri content index",
]
COMMON_SECTION_COMPACTS = [_compact(h) for h in COMMON_SECTION_HEADINGS]

BOUNDARY_SPECS = [
    ("Chairman Message", PRESETS["Chairman Message"]["aliases"]),
    ("CEO Message", PRESETS["CEO Message"]["aliases"]),
    ("Managing Director Message", PRESETS["Managing Director Message"]["aliases"]),
    ("Board's Report", ["board's report", "boards' report", "boards’ report"]),
    ("Directors' Report", ["directors' report", "directors’ report", "director's report", "directors' report to the shareholders"]),
    ("Management Discussion & Analysis", PRESETS["Management Discussion & Analysis"]["aliases"]),
    ("Corporate Governance Report", ["corporate governance report", "report on corporate governance"]),
    ("Corporate Social Responsibility", ["corporate social responsibility", "corporate social responsibility report", "csr report"]),
    ("BRSR", PRESETS["Business Responsibility & Sustainability Report (BRSR)"]["aliases"]),
    ("BRR", PRESETS["Business Responsibility Report (BRR)"]["aliases"]),
    ("ESG Report", PRESETS["ESG Report"]["aliases"]),
    ("Sustainability Report", PRESETS["Sustainability Report"]["aliases"]),
    ("Clinical Governance", ["clinical governance"]),
    ("Standalone Financial Statements", ["standalone financial statements", "standalone financials"]),
    ("Consolidated Financial Statements", ["consolidated financial statements", "consolidated financials"]),
    ("Independent Auditor's Report", ["independent auditor's report", "independent auditors' report", "report on the audit of the standalone financial statements"]),
    ("Notice", ["notice", "notice of annual general meeting"]),
    ("GRI Index", ["gri index", "gri content index"]),
    ("About Company", ["about the company", "about our company", "company overview", "company profile", "corporate profile", "at a glance"]),
    ("Board of Directors", ["board of directors"]),
    ("Management Team", ["management team"]),
    ("The Year Gone By", ["the year gone by"]),
    ("Performance Highlights", ["performance highlights", "financial & operational highlights"]),
    ("Corporate Review", ["corporate review"]),
    ("Awards", ["awards", "awards and accolades"]),
]

_SECTION_MAP_CACHE = {}
_GLOBAL_GRAPH_CACHE = {}
_TOC_CACHE = {}

def _doc_key(pages):
    if not pages:
        return (0,)
    first = pages[0].get("text", "")[:700]
    middle = pages[len(pages) // 2].get("text", "")[:400]
    last = pages[-1].get("text", "")[-700:]
    return (len(pages), hash(first), hash(middle), hash(last))

def _page_lines(page):
    lines = page.get("lines") or []
    if lines: return lines
    return [{"order": i, "text": t, "bbox": None, "size": None} for i, t in enumerate(page.get("text", "").splitlines()) if _norm_line(t)]

def _top_nav_label(page, names=("content", "contents", "index")):
    height = float(page.get("height") or 800.0)
    lines = page.get("native_lines") or _page_lines(page)
    for i, line in enumerate(lines):
        text = _norm_line(line.get("text", "")).lower()
        if text not in names: continue
        bb = line.get("bbox")
        if bb:
            if ((bb[1] + bb[3]) / 2) < 0.35 * height: return True
        elif i < 15:
            return True
    return False

def _toc_structure_score(page):
    cached = page.get("_toc_structure_cache")
    if cached is not None: return tuple(cached)
    text = page.get("text", "")
    lines = [_norm_line(x.get("text", "")) for x in (page.get("native_lines") or _page_lines(page)) if _norm_line(x.get("text", ""))]
    hits = sum(1 for line in lines for h in COMMON_SECTION_COMPACTS if len(_compact(line)) <= 90 and (h == _compact(line) or (h in _compact(line) and len(_compact(line)) <= len(h) + 24)))
    words = len(re.findall(r"[A-Za-z]+", text))
    page["_toc_structure_cache"] = (hits, words)
    return hits, words

def _toc_like_page(page):
    if "_toc_like_cache" in page: return bool(page["_toc_like_cache"])
    value = _top_nav_label(page)
    if not value:
        page_no = page.get("page")
        if page_no is None or page_no <= 15:
            hits, words = _toc_structure_score(page)
            value = hits >= 5 and words <= 500
    page["_toc_like_cache"] = bool(value)
    return bool(value)

def _heading_variants(heading):
    raw = list(heading) if isinstance(heading, (list, tuple, set)) else [heading]
    out, seen = [], set()
    for h in raw:
        if not h: continue
        h = _norm_line(str(h)).replace("’", "'")
        variants = [h, re.sub(r"\s*\([^)]*\)\s*", " ", h).strip()]
        if "&" in h: variants.append(h.replace("&", "and"))
        if re.search(r"\band\b", h, re.I): variants.append(re.sub(r"\band\b", "&", h, flags=re.I))
        for v in variants:
            c = _compact(v)
            if c and c not in seen:
                seen.add(c)
                out.append(v)
    return out

def _window_candidates(page, aliases, max_width=5, lines_override=None, detection_source="layout"):
    aliases = _heading_variants(aliases)
    ac = [(_compact(a), a) for a in aliases]
    lines = lines_override if lines_override is not None else _page_lines(page)
    if not lines: return []

    sizes = [l.get("size") for l in lines if isinstance(l.get("size"), (int, float)) and 5 <= l.get("size") <= 15]
    body = sorted(sizes)[len(sizes) // 2] if sizes else 9.0
    height = float(page.get("height") or 800.0)
    out = []

    for i in range(len(lines)):
        for width in range(1, min(max_width, len(lines) - i) + 1):
            group = lines[i:i + width]
            if any(x.get("layout_class") == "folio-edge" for x in group): continue
            text = " ".join(_norm_line(x.get("text", "")) for x in group).strip()
            wc = _compact(text)
            if not wc: continue
            for target, alias in ac:
                if not target: continue
                exact = (wc == target)
                contained = (target in wc and len(wc) <= len(target) + 30)
                annexure_contained = (target in wc and "annexure" in wc and len(wc) <= len(target) + 45)
                if not (exact or contained or annexure_contained): continue
                if target == _compact("corporate social responsibility") and "committee" in wc: continue

                score = 48 if exact else 42
                
                if exact and target in {"boardsreport", "directorsreport", "directorsreporttotheshareholders"}:
                    score += 15
                
                layout_classes = {str(x.get("layout_class") or "") for x in group}
                if "title" in layout_classes: score += 22
                elif "section-header" in layout_classes: score += 18
                elif "page-header" in layout_classes or "page-footer" in layout_classes: score -= 30
                if annexure_contained: score += 8

                max_size = max((x.get("size") or 0) for x in group)
                delta = max_size - body
                if max_size >= 16: score += 18
                elif delta >= 3: score += 14
                elif delta >= 1.5: score += 8

                bboxes = [x.get("bbox") for x in group if x.get("bbox")]
                if bboxes:
                    x0 = min(b[0] for b in bboxes); y0 = min(b[1] for b in bboxes)
                    x1 = max(b[2] for b in bboxes); y1 = max(b[3] for b in bboxes)
                    if y0 / max(height, 1) < 0.18: score += 8
                    elif y0 / max(height, 1) < 0.35: score += 4
                    bbox = [x0, y0, x1, y1]
                else:
                    bbox = None
                    if i < 12: score += 5

                letters = re.sub(r"[^A-Za-z]+", "", text)
                if letters and text.upper() == text and len(letters) >= 8: score += 3

                out.append({
                    "score": score, "line_order": group[0].get("order", i), "line_end": group[-1].get("order", i + width - 1),
                    "bbox": bbox, "matched_alias": alias, "matched_text": text, "detection_source": detection_source,
                })
    best = {}
    for c in out:
        k = (c["line_order"], c["matched_alias"])
        if k not in best or c["score"] > best[k]["score"]: best[k] = c
    return list(best.values())

def _edge_folios(page):
    cached = page.get("_edge_folios_cache")
    if cached is not None:
        return cached
    width = float(page.get("width") or 600.0); height = float(page.get("height") or 800.0)
    out = []
    for line in (page.get("native_lines") or _page_lines(page)):
        text = _norm_line(line.get("text", "")); bb = line.get("bbox")
        if not bb: continue
        x0, y0, x1, y1 = bb
        strict_edge = y0 < 0.06 * height or y1 > 0.94 * height
        header_footer = y0 < 0.12 * height or y1 > 0.90 * height
        if not header_footer: continue
        nums = []
        m = re.fullmatch(r"0*(\d{1,4})", text)
        if m and strict_edge: nums = [int(m.group(1))]
        elif len(text) <= 150 and header_footer:
            m = re.search(r"(?:\bI\s+|\|\s*|\s)0*(\d{1,4})\s*$", text)
            if m: nums = [int(m.group(1))]
        for n in nums:
            if 0 <= n <= 1500 and not (1900 <= n <= 2099):
                xc = (x0 + x1) / 2
                if width / max(height, 1) >= 1.25 and not (xc <= 0.20 * width or xc >= 0.80 * width): continue
                out.append({"number": n, "x": xc, "y": (y0 + y1) / 2})
    page["_edge_folios_cache"] = out
    return out

def _explicit_toc_page(page):
    if _top_nav_label(page): return True
    hits, words = _toc_structure_score(page)
    return hits >= 5 and words <= 500

def _toc_printed_page(pages, aliases):
    key = (_doc_key(pages), tuple(_compact(x) for x in _heading_variants(aliases)))
    if key in _TOC_CACHE: return _TOC_CACHE[key]
    targets = [_compact(x) for x in _heading_variants(aliases)]
    answer = None
    for page in pages[:30]:
        if not _explicit_toc_page(page): continue
        lines = page.get("native_lines") or _page_lines(page)
        for i, line in enumerate(lines):
            text = _norm_line(line.get("text", "")).replace("’", "'")
            c = _compact(text)
            matched = next((t for t in targets if t and t in c and len(c) <= len(t) + 40), None)
            if not matched: continue
            nums = [int(x) for x in re.findall(r"\b(\d{1,4})\b", text) if not (1900 <= int(x) <= 2099)]
            if nums:
                m1 = re.match(r"^\s*0*(\d{1,4})\b", text); m2 = re.search(r"\b0*(\d{1,4})\s*$", text)
                if m1 and not (1900 <= int(m1.group(1)) <= 2099): answer = int(m1.group(1)); break
                if m2 and not (1900 <= int(m2.group(1)) <= 2099): answer = int(m2.group(1)); break
            bb = line.get("bbox")
            if bb:
                yc = (bb[1] + bb[3]) / 2; numeric = []
                for nline in lines:
                    nt = _norm_line(nline.get("text", "")); nb = nline.get("bbox")
                    mm = re.fullmatch(r"0*(\d{1,4})", nt)
                    if mm and nb and 0 < int(mm.group(1)) <= 1500 and not (1900 <= int(mm.group(1)) <= 2099):
                        ny = (nb[1] + nb[3]) / 2
                        if abs(ny - yc) <= 18: numeric.append((abs(ny - yc), int(mm.group(1)), ny))
                if numeric:
                    preceding = [x for x in numeric if x[2] <= yc + 2] if len(numeric[0]) >= 3 else []
                    answer = min(preceding, key=lambda x: x[0])[1] if preceding else min(numeric, key=lambda x: x[0])[1]
                    break
        if answer is not None: break
    if len(_TOC_CACHE) > 200: _TOC_CACHE.clear()
    _TOC_CACHE[key] = answer
    return answer

def _toc_page_numbers(pages):
    if pages and "_toc_page_numbers_cache" in pages[0]:
        return list(pages[0]["_toc_page_numbers_cache"])
    nums = set()
    skip_labels = {"contents", "content", "index", "table of contents"}
    for page in pages[:30]:
        if not _explicit_toc_page(page): continue
        lines = [x for x in [_norm_line(x.get("text", "")) for x in (page.get("native_lines") or _page_lines(page))] if x]
        for i, t in enumerate(lines):
            if not re.search(r"[A-Za-z]", t) or t.lower().strip(" :.-") in skip_labels: continue
            same = None
            for mm in (re.match(r"^\s*0*(\d{1,4})\b", t), re.search(r"\b0*(\d{1,4})\s*$", t)):
                if mm: same = int(mm.group(1)); break
            if same is not None and 0 < same <= 1500 and not (1900 <= same <= 2099):
                nums.add(same); continue
            if len(t) > 120 or len(t.split()) > 16 or t.endswith(('.', ';')): continue
            if i + 1 < len(lines):
                m = re.fullmatch(r"0*(\d{1,4})", lines[i + 1])
                if m and 0 < int(m.group(1)) <= 1500 and not (1900 <= int(m.group(1)) <= 2099): nums.add(int(m.group(1)))
    result = sorted(nums)
    if pages:
        pages[0]["_toc_page_numbers_cache"] = tuple(result)
    return result

def _is_two_up_document(pages):
    if pages and "_two_up_document_cache" in pages[0]:
        return bool(pages[0]["_two_up_document_cache"])
    hits = 0
    for page in pages[:120]:
        w = float(page.get("width") or 0); h = float(page.get("height") or 0)
        if w and h and w / max(h, 1) >= 1.25 and len({f["number"] for f in _edge_folios(page)}) >= 2:
            hits += 1
            if hits >= 2:
                if pages:
                    pages[0]["_two_up_document_cache"] = True
                return True
    if pages:
        pages[0]["_two_up_document_cache"] = False
    return False

def _anchor_for_printed_page(pages, printed_page, min_index=0):
    if printed_page is None: return None
    for idx, page in enumerate(pages):
        if idx < min_index: continue
        matches = [f for f in _edge_folios(page) if f["number"] == printed_page]
        if matches:
            x = float(matches[0].get("x") or 0)
            return {"label": "TOC logical boundary", "index": idx, "pdf_page": page.get("page"), "line_order": 0, "line_end": 0, "bbox": [x - 1.0, 0.0, x + 1.0, 1.0], "printed_page": printed_page, "score": 100, "detection_source": "toc-logical-boundary"}
    return None

def _anchor_is_two_up(pages, anchor):
    if not anchor or anchor.get("index") is None or not (0 <= anchor["index"] < len(pages)): return False
    page = pages[anchor["index"]]
    w = float(page.get("width") or 0); h = float(page.get("height") or 0)
    return w and h and w / max(h, 1) >= 1.25 and len({f["number"] for f in _edge_folios(page)}) >= 2

def _next_toc_logical_boundary(pages, start_anchor):
    pp = start_anchor.get("printed_page")
    if pp is None or not _anchor_is_two_up(pages, start_anchor): return None
    later = [n for n in _toc_page_numbers(pages) if n > pp]
    return _anchor_for_printed_page(pages, min(later), min_index=start_anchor.get("index", 0)) if later else None

def _separate_enclosure_status(pages, aliases):
    compact_aliases = [_compact(a) for a in _heading_variants(aliases)]
    rows = []
    for page in pages[:35]:
        rows.extend(_norm_line(x) for x in page.get("text", "").splitlines() if _norm_line(x))
        rows.extend(_norm_line(x.get("text", "")) for x in (page.get("native_lines") or []) if _norm_line(x.get("text", "")))
    for i, line in enumerate(rows):
        window = " ".join(rows[max(0, i - 1): min(len(rows), i + 3)])
        if any(a and a in _compact(window) for a in compact_aliases):
            low = window.lower().replace("’", "'")
            if re.search(r"web\s*[- ]?link\s+only|available\s+only\s+(?:on|at)\s+(?:the\s+)?website|website\s+only", low): return "web-only"
            if re.search(r"separate\s+(?:enclosure|booklet|report)|separately\s+(?:enclosed|attached|available)|enclosed\s+separately|attached\s+separately", low): return "separate"
    return None

def _find_separate_appendix_anchor(pages, aliases):
    if not pages: return None
    graph = build_global_section_graph(pages)
    financial_idxs = [n.get("index") for n in graph if _canonical_request_label(n.get("label")) in {"Standalone Financial Statements", "Consolidated Financial Statements"} and isinstance(n.get("index"), int)]
    tail_start = min(len(pages), max(int(len(pages) * 0.55), (max(financial_idxs) + 1) if financial_idxs else 0))
    sliced = pages[tail_start:]
    if not sliced: return None
    a = _detect_section_anchor(sliced, aliases, min_score=60)
    if not a: return None
    a = dict(a); a["index"] = a.get("index", 0) + tail_start
    a["pdf_page"] = pages[a["index"]].get("page") if 0 <= a["index"] < len(pages) else a.get("pdf_page")
    a["detection_source"] = "separate-appendix"
    return a

def _anchor_printed_page(page, anchor, toc_page=None):
    if toc_page is not None: return toc_page
    folios = _edge_folios(page)
    if not folios: return None
    height = float(page.get("height") or 800.0)
    best_edge = min(min(f["y"], max(0.0, height - f["y"])) for f in folios)
    edge_folios = [f for f in folios if min(f["y"], max(0.0, height - f["y"])) <= best_edge + 12]
    if anchor.get("bbox"):
        x = (anchor["bbox"][0] + anchor["bbox"][2]) / 2
        return min(edge_folios, key=lambda f: abs(f["x"] - x))["number"]
    return edge_folios[0]["number"]

def _detect_section_anchor(pages, aliases, min_score=54):
    aliases = _heading_variants(aliases)
    toc_page = _toc_printed_page(pages, aliases)
    best = None
    target_compacts = [_compact(a) for a in aliases if _compact(a)]

    for idx, page in enumerate(pages):
        if _toc_like_page(page): continue
        page_compact = page.setdefault("_page_compact_cache", _compact(page.get("text", "")))
        native_compact = page.setdefault("_native_compact_cache", _compact(" ".join(x.get("text", "") for x in (page.get("native_lines") or []))))
        ocr_compact = page.setdefault("_ocr_compact_cache", _compact(" ".join(x.get("text", "") for x in (page.get("ocr_heading_lines") or []))))
        if not any(t in page_compact or t in native_compact or t in ocr_compact for t in target_compacts): continue

        candidate_sets = [(_page_lines(page), "layout")]
        if page.get("native_lines"): candidate_sets.append((page.get("native_lines"), "native"))
        if page.get("ocr_heading_lines"): candidate_sets.append((page.get("ocr_heading_lines"), "ocr-probe"))

        seen = set()
        for det_lines, source in candidate_sets:
            for cand in _window_candidates(page, aliases, lines_override=det_lines, detection_source=source):
                ck = (_compact(cand.get("matched_text", "")), tuple(round(x, 1) for x in (cand.get("bbox") or [])))
                if ck in seen and source == "native": continue
                seen.add(ck)

                item = dict(cand); item.update({"index": idx, "pdf_page": page.get("page")})
                actual_pp = _anchor_printed_page(page, item)
                score = cand["score"]

                if toc_page is not None and actual_pp is not None:
                    dist = abs(actual_pp - toc_page)
                    if dist == 0: score += 45
                    elif dist == 1: score += 12
                    elif dist > 3: score -= 15

                if source == "native": score += 2
                elif source == "ocr-probe": score += 18

                item["score"] = score; item["printed_page"] = actual_pp
                if actual_pp is None and toc_page is not None and score >= 66: item["printed_page"] = toc_page

                distance_rank = -abs((actual_pp if actual_pp is not None else 10**6) - toc_page) if toc_page is not None else 0
                rank = (score, distance_rank, -idx, -item.get("line_order", 0))
                if best is None or rank > best[0]: best = (rank, item)

    if not best or best[1]["score"] < min_score: return None
    return best[1]

def _native_exact_heading_anchor(pages, aliases):
    aliases = _heading_variants(aliases)
    targets = [(_compact(a), a) for a in aliases if _compact(a)]
    toc_page = _toc_printed_page(pages, aliases)
    best = None
    for idx, page in enumerate(pages):
        if _toc_like_page(page): continue
        for j, line in enumerate(page.get("native_lines") or []):
            text = _norm_line(line.get("text", "")); tc = _compact(text)
            if not tc or len(text) > 150: continue
            matched = next((alias for target, alias in targets if tc == target or (target in tc and len(tc) <= len(target) + 18)), None)
            if not matched: continue
            anchor = {"score": 76, "line_order": line.get("order", j), "line_end": line.get("order", j), "bbox": line.get("bbox"), "matched_alias": matched, "matched_text": text, "detection_source": "native-exact-fallback", "index": idx, "pdf_page": page.get("page")}
            pp = _anchor_printed_page(page, anchor); anchor["printed_page"] = pp
            dist = 9999
            if toc_page is not None and pp is not None:
                dist = abs(pp - toc_page)
                if dist == 0: anchor["score"] += 42
                elif dist == 1: anchor["score"] += 10
                elif dist > 3: anchor["score"] -= 20
            rank = (anchor["score"], -dist, -idx, -anchor["line_order"])
            if best is None or rank > best[0]: best = (rank, anchor)
    return best[1] if best and best[1]["score"] >= 70 else None

def _native_visual_wrapper_anchor(pages, aliases, kind="report"):
    aliases = _heading_variants(aliases)
    targets = [(_compact(a), a) for a in aliases if _compact(a)]
    best = None
    for idx, page in enumerate(pages):
        lines = page.get("native_lines") or []
        if not lines:
            continue
        for j, line in enumerate(lines):
            text = _norm_line(line.get("text", "")).replace("’", "'")
            tc = _compact(text)
            matched = next((a for t, a in targets if tc == t or (t in tc and len(tc) <= len(t) + 10)), None)
            if not matched or len(text) > 100:
                continue
            following = " ".join(_norm_line(x.get("text", "")) for x in lines[j + 1:j + 65]).lower().replace("’", "'")
            if kind == "board":
                signals = sum(bool(re.search(pat, following)) for pat in (
                    r"\bto the members\b", r"\bannual report\b", r"\bboard of directors\b",
                    r"\bfinancial results?\b", r"\byear ended\b",
                ))
                need = 2
            else:
                signals = sum(bool(re.search(pat, following)) for pat in (
                    r"\bsection a\b", r"general (?:information|disclosures)", r"corporate identity number",
                    r"name of the (?:company|listed entity)", r"registered (?:office )?address",
                    r"financial year (?:reported|for which reporting is being done)",
                ))
                need = 2
            if signals < need:
                continue
            anchor = {
                "score": 132 + min(signals, 5), "line_order": line.get("order", j),
                "line_end": line.get("order", j), "bbox": line.get("bbox"),
                "matched_alias": matched, "matched_text": _norm_line(line.get("text", "")),
                "detection_source": f"native-{kind}-wrapper", "index": idx, "pdf_page": page.get("page"),
            }
            anchor["printed_page"] = _anchor_printed_page(page, anchor)
            rank = (anchor["score"], -idx, -anchor["line_order"])
            if best is None or rank > best[0]:
                best = (rank, anchor)
    return best[1] if best else None


def _discover_annexure_anchors(pages):
    out = []
    pat = re.compile(r"^\s*annexure\s*[-:]?\s*([ivxlcdm]+|\d{1,2})\b", re.I)
    for idx, page in enumerate(pages):
        if _toc_like_page(page): continue
        height = float(page.get("height") or 800.0)
        lines = _page_lines(page)
        sizes = [l.get("size") for l in lines if isinstance(l.get("size"), (int, float)) and 5 <= l.get("size") <= 15]
        body = sorted(sizes)[len(sizes)//2] if sizes else 9.0
        for line in lines[:40]:
            text = _norm_line(line.get("text", ""))
            m = pat.match(text)
            if not m: continue
            bb = line.get("bbox"); size = line.get("size") or body
            if bb and bb[1] > 0.30 * height and size < body + 1.5: continue
            anchor = {"label": f"Annexure {m.group(1)}", "aliases": [text], "index": idx, "pdf_page": page.get("page"), "line_order": line.get("order", 0), "line_end": line.get("order", 0), "bbox": bb, "score": 70, "matched_alias": text, "matched_text": text}
            anchor["printed_page"] = _anchor_printed_page(page, anchor)
            out.append(anchor)
            break
    return out

ANNEXURE_CANONICAL_SPECS = [
    ("Management Discussion & Analysis", PRESETS["Management Discussion & Analysis"]["aliases"]),
    ("Corporate Governance Report", ["corporate governance report", "report on corporate governance"]),
    ("Corporate Social Responsibility", ["corporate social responsibility report", "corporate social responsibility", "csr report"]),
    ("BRSR", PRESETS["Business Responsibility & Sustainability Report (BRSR)"]["aliases"]),
    ("BRR", PRESETS["Business Responsibility Report (BRR)"]["aliases"]),
    ("ESG Report", PRESETS["ESG Report"]["aliases"]),
    ("Sustainability Report", PRESETS["Sustainability Report"]["aliases"]),
]

def _annexure_toc_map(pages):
    result = {}
    annex_pat = re.compile(r"^\s*annexure\s*[-:]?\s*([ivxlcdm]+|\d{1,3}|[a-z])\b\s*(.*)$", re.I)
    for page in pages[:30]:
        if not _explicit_toc_page(page): continue
        lines = [x for x in [_norm_line(x.get("text", "")) for x in (page.get("native_lines") or _page_lines(page))] if x]
        for i, text in enumerate(lines):
            m = annex_pat.match(text.replace("’", "'"))
            if not m: continue
            key = f"annexure {_compact(m.group(1))}"
            remainder = re.sub(r"^\s*[-:–—]+\s*", "", m.group(2) or "").strip()
            remainder = re.sub(r"^\d{1,4}\s+|\s+\d{1,4}$", "", remainder).strip()
            title = remainder if re.search(r"[A-Za-z]", remainder) else ""
            if not title:
                for j in range(i + 1, min(len(lines), i + 4)):
                    nxt = lines[j]
                    if annex_pat.match(nxt) or re.fullmatch(r"0*\d{1,4}", nxt): continue
                    if re.search(r"[A-Za-z]", nxt) and len(nxt) <= 140:
                        title = re.sub(r"^\d{1,4}\s+|\s+\d{1,4}$", "", nxt).strip(); break
            if title: result[key] = title
    return result

def _canonical_from_semantic_text(text):
    wc = _compact(text)
    for label, aliases in ANNEXURE_CANONICAL_SPECS:
        for alias in _heading_variants(aliases):
            ac = _compact(alias)
            if ac and ac in wc: return label, alias
    return None, None

def resolve_annexure_headings(pages):
    annex_pat = re.compile(r"^\s*annexure\s*[-:]?\s*([ivxlcdm]+|\d{1,3}|[a-z])\b", re.I)
    toc_map = _annexure_toc_map(pages)
    out = []
    for idx, page in enumerate(pages):
        if _toc_like_page(page): continue
        lines = page.get("native_lines") or _page_lines(page)
        if not lines: continue
        for i, line in enumerate(lines[:40]):
            txt = _norm_line(line.get("text", "")).replace("’", "'")
            m = annex_pat.match(txt)
            if not m: continue
            group = lines[i:min(len(lines), i + 9)]
            window = " ".join(_norm_line(x.get("text", "")) for x in group)
            annex_key = f"annexure {_compact(m.group(1))}"
            semantic = window + (" " + toc_map[annex_key] if toc_map.get(annex_key) else "")
            label, matched_alias = _canonical_from_semantic_text(semantic)
            bboxes = [x.get("bbox") for x in group if x.get("bbox")]
            bbox = [min(b[0] for b in bboxes), min(b[1] for b in bboxes), max(b[2] for b in bboxes), max(b[3] for b in bboxes)] if bboxes else None
            
            if label:
                anchor = {"label": label, "aliases": [matched_alias], "index": idx, "pdf_page": page.get("page"), "line_order": line.get("order", i), "line_end": group[-1].get("order", i + len(group) - 1), "bbox": bbox, "score": 98, "matched_alias": matched_alias, "matched_text": window, "annexure_key": f"Annexure {m.group(1)}", "detection_source": "annexure-semantic"}
                anchor["printed_page"] = _anchor_printed_page(page, anchor)
                out.append(anchor)
            else:
                local_title = re.sub(annex_pat, "", txt).strip(" :-–—") or toc_map.get(annex_key, "")
                if local_title and _looks_like_heading_text(local_title):
                    anchor = {"label": "Generic Boundary", "aliases": [local_title], "index": idx, "pdf_page": page.get("page"), "line_order": line.get("order", i), "line_end": group[-1].get("order", i + len(group) - 1), "bbox": bbox, "score": 82, "matched_alias": local_title, "matched_text": window, "annexure_key": f"Annexure {m.group(1)}", "generic_boundary": True, "detection_source": "annexure-generic"}
                    anchor["printed_page"] = _anchor_printed_page(page, anchor)
                    out.append(anchor)
            break
    return out

def _anchor_key(a): return (a["index"], a.get("line_order", 0))

def _median(values, default=0.0):
    vals = sorted(float(v) for v in values if isinstance(v, (int, float)))
    if not vals: return float(default)
    n = len(vals)
    return vals[n // 2] if n % 2 else (vals[n // 2 - 1] + vals[n // 2]) / 2.0

def _looks_like_heading_text(text):
    text = _norm_line(text)
    if not (2 <= len(text) <= 150) or not re.search(r"[A-Za-z]", text): return False
    words = re.findall(r"[A-Za-z][A-Za-z&'’/-]*", text)
    if not (1 <= len(words) <= 18) or text.endswith((".", ";", ",")): return False
    low = text.lower().strip(" :-–—")
    noisy_prefixes = ("section a", "section b", "section c", "principle ", "note ", "notes ", "sr no", "sr. no", "particulars", "question ", "table ", "figure ", "amount in", "(`", "(%", "source:", "source ", "page ")
    if any(low.startswith(x) for x in noisy_prefixes): return False
    if sum(ch.isdigit() for ch in text) > max(6, len(text) * 0.25): return False
    return True

def _doc_body_font_size(pages):
    if pages and "_doc_body_font_size_cache" in pages[0]:
        return float(pages[0]["_doc_body_font_size_cache"])
    vals = [line.get("size") for page in pages[: min(len(pages), 180)] for line in (page.get("native_lines") or []) if isinstance(line.get("size"), (int, float)) and 6 <= line.get("size") <= 14 and len(_norm_line(line.get("text", ""))) >= 30]
    result = _median(vals, 9.0)
    if pages:
        pages[0]["_doc_body_font_size_cache"] = result
    return result

def _anchor_native_font_size(page, anchor):
    target = _compact(anchor.get("matched_text", ""))
    best = 0.0; abb = anchor.get("bbox")
    for line in (page.get("native_lines") or []):
        sz = line.get("size")
        if not isinstance(sz, (int, float)): continue
        txt = _compact(line.get("text", ""))
        related = bool(target and txt and (txt in target or target in txt))
        bb = line.get("bbox")
        overlaps = bool(abb and bb and not (bb[2] < abb[0] or abb[2] < bb[0] or bb[3] < abb[1] or abb[3] < bb[1]))
        if related or overlaps: best = max(best, float(sz))
    return best

def _primary_heading_font_size(pages, seed_nodes):
    vals = [sz for a in seed_nodes if isinstance(a.get("index"), int) and 0 <= a["index"] < len(pages) and not str(a.get("label", "")).lower().startswith("annexure") and (sz := _anchor_native_font_size(pages[a["index"]], a))]
    body = _doc_body_font_size(pages)
    return max(body + 1.8, _median(vals, body + 2.5)) if vals else body + 2.8

def _generic_toc_entries(pages):
    if pages and "_generic_toc_entries_cache" in pages[0]:
        return [dict(x) for x in pages[0]["_generic_toc_entries_cache"]]
    entries = []
    skip = {"content", "contents", "index", "table of contents"}
    for page in pages[:30]:
        if not _explicit_toc_page(page): continue
        lines = page.get("native_lines") or _page_lines(page)
        for i, line in enumerate(lines):
            text = _norm_line(line.get("text", ""))
            if not _looks_like_heading_text(text): continue
            clean = text.strip(); low = clean.lower().strip(" :-–—")
            if low in skip: continue
            number, title = None, clean
            m_left = re.match(r"^\s*0*(\d{1,4})\s+(.+)$", clean); m_right = re.match(r"^(.+?)\s+0*(\d{1,4})\s*$", clean)
            if m_left and not (1900 <= int(m_left.group(1)) <= 2099): number = int(m_left.group(1)); title = m_left.group(2).strip()
            elif m_right and not (1900 <= int(m_right.group(2)) <= 2099): number = int(m_right.group(2)); title = m_right.group(1).strip()
            if number is None:
                bb = line.get("bbox")
                if bb:
                    yc = (bb[1] + bb[3]) / 2; nums = []
                    for nline in lines:
                        nt = _norm_line(nline.get("text", "")); nb = nline.get("bbox")
                        mm = re.fullmatch(r"0*(\d{1,4})", nt)
                        if mm and nb and 0 < int(mm.group(1)) <= 1500 and not (1900 <= int(mm.group(1)) <= 2099):
                            ny = (nb[1] + nb[3]) / 2
                            if abs(ny - yc) <= 18: nums.append((abs(ny - yc), int(mm.group(1))))
                    if nums: number = min(nums)[1]
            title = re.sub(r"^\s*[-–—:]+|[-–—:]+\s*$", "", title).strip()
            if number is not None and _looks_like_heading_text(title):
                entries.append({"title": title, "compact": _compact(title), "printed_page": number})
    best = {}
    for e in entries:
        k = (e["compact"], e["printed_page"])
        if k not in best or len(e["title"]) > len(best[k]["title"]): best[k] = e
    result = list(best.values())
    if pages:
        pages[0]["_generic_toc_entries_cache"] = tuple(dict(x) for x in result)
    return result

def _toc_match_for_heading(text, printed_page, entries):
    tc = _compact(text)
    if not tc: return None
    best = None
    for e in entries:
        ec = e["compact"]
        if not ec: continue
        exactish = tc == ec or (tc in ec and len(ec) <= len(tc) + 26) or (ec in tc and len(tc) <= len(ec) + 26)
        ratio = SequenceMatcher(None, tc, ec).ratio() if not exactish else 1.0
        if not exactish and ratio < 0.88: continue
        pp_score = 0
        if printed_page is not None:
            d = abs(int(printed_page) - int(e["printed_page"]))
            if d == 0: pp_score = 3
            elif d == 1: pp_score = 1
            elif d > 3: continue
        rank = (pp_score, ratio, len(ec))
        if best is None or rank > best[0]: best = (rank, e)
    return best[1] if best else None

def _layout_class_for_native_line(page, text, bbox=None):
    tc = _compact(text)
    cached = page.get("_layout_heading_rows_cache")
    if cached is None:
        cached = [(lc, line.get("bbox"), cls) for line in _page_lines(page) if (cls := str(line.get("layout_class") or "")) in {"title", "section-header"} and (lc := _compact(line.get("text", "")))]
        page["_layout_heading_rows_cache"] = cached
    best = None
    for lc, lbb, cls in cached:
        related = tc == lc or tc in lc or lc in tc
        if not related and bbox and lbb: related = not (lbb[2] < bbox[0] or bbox[2] < lbb[0] or lbb[3] < bbox[1] or bbox[3] < lbb[1])
        if related and best != "title": best = cls
    return best

def _discover_generic_top_level_anchors(pages, seed_nodes):
    body = _doc_body_font_size(pages)
    primary = _primary_heading_font_size(pages, seed_nodes)
    toc_entries = _generic_toc_entries(pages)
    seed_positions = {(a.get("index"), a.get("line_order", 0)) for a in seed_nodes}
    out = []

    for idx, page in enumerate(pages):
        if _toc_like_page(page): continue
        lines = page.get("native_lines") or []
        if not lines: continue
        height = float(page.get("height") or 800.0)
        for i, line in enumerate(lines):
            if (idx, line.get("order", i)) in seed_positions: continue
            text = _norm_line(line.get("text", ""))
            if not _looks_like_heading_text(text): continue
            bb = line.get("bbox")
            if bb and (bb[1] < 0.07 * height or bb[3] > 0.93 * height): continue
            sz = float(line.get("size") or 0.0); bold = bool(line.get("bold"))
            pp_probe = {"bbox": bb, "line_order": line.get("order", i), "matched_text": text}
            printed_page = _anchor_printed_page(page, pp_probe)
            toc_match = _toc_match_for_heading(text, printed_page, toc_entries)
            layout_cls = _layout_class_for_native_line(page, text, bb)

            prev_bb = lines[i - 1].get("bbox") if i > 0 else None
            next_bb = lines[i + 1].get("bbox") if i + 1 < len(lines) else None
            gap_before = (bb[1] - prev_bb[3]) if bb and prev_bb else (20.0 if i == 0 else 0.0)
            gap_after = (next_bb[1] - bb[3]) if bb and next_bb else 10.0
            isolated = gap_before >= max(6.0, body * 0.65) or gap_after >= max(6.0, body * 0.65)
            strongly_isolated = gap_before >= max(10.0, body) and gap_after >= max(6.0, body * 0.65)
            near_top = bool(bb and bb[1] <= 0.30 * height)

            score = 0
            if toc_match:
                score += 48
                if printed_page is not None and printed_page == toc_match["printed_page"]: score += 12
            if layout_cls == "title": score += 34
            elif layout_cls == "section-header": score += 28
            if sz >= primary * 1.05: score += 34
            elif sz >= primary * 0.95: score += 26
            elif sz >= body + 2.0: score += 16
            if bold: score += 8
            if isolated: score += 8
            if strongly_isolated: score += 7
            if near_top: score += 6
            letters = re.sub(r"[^A-Za-z]", "", text)
            if letters and text.upper() == text and len(letters) >= 8: score += 4

            accept = False
            if toc_match and score >= 62: accept = True
            elif score >= 84 and sz >= primary * 0.95 and (strongly_isolated or layout_cls in {"title", "section-header"}): accept = True
            if not accept: continue

            hard_boundary = bool(toc_match and sz >= primary * 0.95)
            
            anchor = {
                "label": "Generic Boundary", "aliases": [text], "index": idx, "pdf_page": page.get("page"),
                "line_order": line.get("order", i), "line_end": line.get("order", i), "bbox": bb,
                "score": score, "matched_alias": text, "matched_text": text,
                "printed_page": printed_page if printed_page is not None else (toc_match["printed_page"] if toc_match else None),
                "generic_boundary": True, "toc_supported": bool(toc_match),
                "hard_boundary": hard_boundary, "soft_boundary": not hard_boundary,
                "detection_source": "generic-visual", "primary_heading_font": primary, "body_font": body,
            }
            out.append(anchor)
    return out

def _same_anchor_region(a, b):
    if a.get("index") != b.get("index"): return False
    ao = a.get("line_order", 0); bo = b.get("line_order", 0)
    if abs(ao - bo) <= 1: return True
    ab = a.get("bbox"); bb = b.get("bbox")
    if ab and bb: return not (ab[2] < bb[0] or bb[2] < ab[0] or ab[3] < bb[1] or bb[3] < ab[1])
    return False

def _merge_global_nodes(nodes):
    merged = []
    for item in sorted(nodes, key=_anchor_key):
        duplicate_at = None
        for j, old in enumerate(merged):
            if not _same_anchor_region(item, old): continue
            ic = _compact(item.get("matched_text", "")); oc = _compact(old.get("matched_text", ""))
            related = not ic or not oc or ic in oc or oc in ic or SequenceMatcher(None, ic, oc).ratio() >= 0.82
            if related: duplicate_at = j; break
        if duplicate_at is None: merged.append(item); continue
        old = merged[duplicate_at]
        rank_item = (0 if item.get("generic_boundary") else 1, 1 if item.get("custom_requested") else 0, item.get("score", 0), len(_compact(item.get("matched_text", ""))))
        rank_old = (0 if old.get("generic_boundary") else 1, 1 if old.get("custom_requested") else 0, old.get("score", 0), len(_compact(old.get("matched_text", ""))))
        if rank_item > rank_old: merged[duplicate_at] = item
    return sorted(merged, key=_anchor_key)

def build_section_map(pages):
    key = _doc_key(pages)
    if key in _SECTION_MAP_CACHE: return _SECTION_MAP_CACHE[key]

    found = []
    for label, aliases in BOUNDARY_SPECS:
        min_score = 64 if label in {"Corporate Social Responsibility"} else 54
        a = _detect_section_anchor(pages, aliases, min_score=min_score)
        if a: a = dict(a); a["label"] = label; a["aliases"] = aliases; found.append(a)
    found.extend(_discover_annexure_anchors(pages))
    found.extend(resolve_annexure_headings(pages))

    merged = []
    for item in sorted(found, key=_anchor_key):
        replaced = False
        for j, old in enumerate(merged):
            if item["index"] != old["index"]: continue
            a0, a1 = item.get("line_order", 0), item.get("line_end", item.get("line_order", 0))
            b0, b1 = old.get("line_order", 0), old.get("line_end", old.get("line_order", 0))
            overlap = not (a1 < b0 - 1 or b1 < a0 - 1)
            mt = _compact(item.get("matched_text", "")); ot = _compact(old.get("matched_text", ""))
            related = mt in ot or ot in mt
            if overlap and related:
                rank_item = (item.get("score", 0), len(mt), -a0); rank_old = (old.get("score", 0), len(ot), -b0)
                if rank_item > rank_old: merged[j] = item
                replaced = True; break
        if not replaced: merged.append(item)
    result = sorted(merged, key=_anchor_key)

    if len(_SECTION_MAP_CACHE) >= 24: _SECTION_MAP_CACHE.clear()
    _SECTION_MAP_CACHE[key] = result
    return result

def _canonical_request_label(label):
    return {"Business Responsibility & Sustainability Report (BRSR)": "BRSR", "Business Responsibility Report (BRR)": "BRR", "Management Discussion & Analysis": "Management Discussion & Analysis", "Sustainability Report": "Sustainability Report", "ESG Report": "ESG Report", "Chairman Message": "Chairman Message", "CEO Message": "CEO Message", "Managing Director Message": "Managing Director Message"}.get(label, label)

_FRONT_MATTER_LABELS = {
    "About Company", "Board of Directors", "Management Team", "The Year Gone By",
    "Performance Highlights", "Corporate Review", "Awards",
}
_STATUTORY_PARENT_LABELS = {
    "Board's Report", "Directors' Report", "Management Discussion & Analysis",
    "Corporate Governance Report", "Corporate Social Responsibility", "BRSR", "BRR",
    "ESG Report", "Sustainability Report", "Clinical Governance",
}

def build_global_section_graph(pages, custom_headings=None):
    doc_key = _doc_key(pages)
    base_nodes = None
    if doc_key in _GLOBAL_GRAPH_CACHE: base_nodes = [dict(x) for x in _GLOBAL_GRAPH_CACHE[doc_key]]

    if base_nodes is None:
        nodes = [dict(x) for x in build_section_map(pages)]
        nodes.extend(_discover_generic_top_level_anchors(pages, nodes))
        nodes = _merge_global_nodes(nodes)

        best = {}
        for node in nodes:
            label = _canonical_request_label(node.get("label"))
            pos = (node.get("index"), node.get("line_order", 0), node.get("printed_page"))
            key = (pos, label)
            rank = (1 if node.get("annexure_key") else 0, 0 if node.get("generic_boundary") else 1, node.get("score", 0), len(_compact(node.get("matched_text", ""))))
            if key not in best or rank > best[key][0]: best[key] = (rank, node)
        base_nodes = sorted([v[1] for v in best.values()], key=_anchor_key)
        if len(_GLOBAL_GRAPH_CACHE) >= 24: _GLOBAL_GRAPH_CACHE.clear()
        _GLOBAL_GRAPH_CACHE[doc_key] = [dict(x) for x in base_nodes]

    if not custom_headings: return [dict(x) for x in base_nodes]

    nodes = [dict(x) for x in base_nodes]
    for heading in custom_headings:
        heading = _norm_line(heading)
        if not heading: continue
        a = _detect_section_anchor(pages, [heading])
        if a: a = dict(a); a.update({"label": heading, "aliases": [heading], "custom_requested": True, "detection_source": a.get("detection_source") or "custom-visual"}); nodes.append(a)
    return _merge_global_nodes(nodes)

def _candidate_local_context(pages, candidate, before_lines=8, after_lines=28):
    """Return local body context around a candidate heading without company/page assumptions."""
    idx = candidate.get("index")
    if not isinstance(idx, int) or not (0 <= idx < len(pages)):
        return ""
    lines = pages[idx].get("native_lines") or _page_lines(pages[idx])
    if not lines:
        return _norm_line(pages[idx].get("text", ""))
    order = int(candidate.get("line_order", 0) or 0)
    # line order is normally sequential, but locate the nearest row defensively.
    pos = min(range(len(lines)), key=lambda i: abs(int(lines[i].get("order", i) or i) - order))
    lo = max(0, pos - before_lines)
    hi = min(len(lines), pos + after_lines + 1)
    return " ".join(_norm_line(x.get("text", "")) for x in lines[lo:hi] if _norm_line(x.get("text", "")))


def _candidate_has_top_level_evidence(pages, candidate):
    """Generic peer-boundary evidence gate.

    TOC agreement, explicit annexure semantics / logical boundaries, or strong visual
    prominence can establish a real document-level section. A large TOC mismatch is
    negative evidence, preventing in-body cross-references from becoming peers.
    """
    source = str(candidate.get("detection_source") or "")
    if source in {"annexure-semantic", "separate-appendix", "toc-logical-boundary"}:
        return True

    aliases = candidate.get("aliases") or [candidate.get("matched_alias"), candidate.get("matched_text")]
    aliases = [a for a in aliases if a]
    toc_page = _toc_printed_page(pages, aliases) if aliases else None
    printed = candidate.get("printed_page")
    if toc_page is not None and printed is not None:
        dist = abs(int(toc_page) - int(printed))
        if dist <= 1:
            return True
        if dist > 3:
            return False

    idx = candidate.get("index")
    if isinstance(idx, int) and 0 <= idx < len(pages):
        sz = _anchor_native_font_size(pages[idx], candidate)
        body = _doc_body_font_size(pages)
        layout_cls = _layout_class_for_native_line(
            pages[idx], candidate.get("matched_text", ""), candidate.get("bbox")
        )
        # Strong visual evidence is allowed when no reliable TOC evidence exists.
        if sz and sz >= body + 2.0 and candidate.get("score", 0) >= 68:
            return True
        if layout_cls == "title" and candidate.get("score", 0) >= 72:
            return True
    return False


_BOARD_REFERENCE_LABELS = {
    "Management Discussion & Analysis", "Corporate Governance Report",
    "Corporate Social Responsibility", "BRSR", "BRR", "ESG Report",
    "Sustainability Report", "Independent Auditor's Report",
}

_BOARD_MAJOR_PEER_LABELS = {
    "Management Discussion & Analysis", "Corporate Governance Report",
    "BRSR", "BRR", "ESG Report", "Sustainability Report",
    "Independent Auditor's Report",
}

_BOARD_ATTACHMENT_TERMS = (
    "secretarial audit", "annual return", "mgt-9", "mgt 9", "form no", "form mr",
    "remuneration of directors", "conservation of energy", "technology absorption",
    "foreign exchange earnings", "particulars of employees", "particulars of remuneration",
)

_ATTACHMENT_CONTEXT_RE = re.compile(
    r"(?i)\b(?:annex(?:ure|ed)|annexed\s+herewith|attached|appended|enclosed|"
    r"forms?\s+(?:a\s+)?part\s+of|forming\s+part\s+of|set\s+out\s+in|"
    r"given\s+in|provided\s+in|refer(?:red)?\s+to\s+as)\b"
)


def _boundary_is_child(parent_label, candidate, pages):
    parent = _canonical_request_label(parent_label)
    child = _canonical_request_label(candidate.get("label"))
    ctext = _norm_line(candidate.get("matched_text", "")).lower().replace("’", "'")

    if parent in {"Board's Report", "Directors' Report"}:
        source = str(candidate.get("detection_source") or "")

        # Plain Annexure N wrappers are part of the parent report. A semantically resolved
        # major section (e.g. an Annexure titled MDA) may still be a true next peer.
        if str(candidate.get("label", "")).lower().startswith("annexure"):
            return True
        if candidate.get("annexure_key") and child not in _BOARD_MAJOR_PEER_LABELS:
            return True
        if candidate.get("generic_boundary") and "annexure" in ctext and source != "annexure-semantic":
            return True

        # Common statutory attachments are normally Board-report children unless an actual
        # independently resolved major report node is present.
        if any(term in ctext for term in _BOARD_ATTACHMENT_TERMS) and child not in _BOARD_MAJOR_PEER_LABELS:
            return True

        if child in _BOARD_REFERENCE_LABELS:
            # Actual semantic annexures / appendix sections are real section starts.
            if source in {"annexure-semantic", "separate-appendix", "toc-logical-boundary"}:
                return False

            local = _candidate_local_context(pages, candidate).lower().replace("’", "'")
            # Board reports frequently contain headings that merely introduce a paragraph
            # saying the actual report is annexed/attached later. Such mentions stay children.
            if _ATTACHMENT_CONTEXT_RE.search(local):
                return True

            # If this occurrence does not have independent top-level evidence, keep it inside
            # the Board/Directors report instead of truncating the parent.
            if not _candidate_has_top_level_evidence(pages, candidate):
                return True

    if child == parent:
        return True

    if parent in _STATUTORY_PARENT_LABELS and child in _FRONT_MATTER_LABELS:
        return True

    if parent in {"BRSR", "BRR", "ESG Report", "Sustainability Report"}:
        assurance_terms = (
            "assurance statement", "assurance report", "limited assurance", "reasonable assurance",
            "verification statement", "verification report", "independent assurance", "assurance opinion",
        )
        if child == "Independent Auditor's Report" or any(t in ctext for t in assurance_terms):
            idx = candidate.get("index")
            if isinstance(idx, int) and 0 <= idx < len(pages):
                window = " ".join(
                    pages[j].get("text", "")
                    for j in range(max(0, idx - 1), min(len(pages), idx + 2))
                ).lower()
                if any(k in window for k in (
                    "assurance", "sustainability", "business responsibility", "brsr", "esg", "responsibility report"
                )):
                    return True
    return False

def _next_boundary_anchor(pages, start_anchor, requested_label=None, custom_headings=None):
    toc_boundary = _next_toc_logical_boundary(pages, start_anchor)
    if toc_boundary is not None: return toc_boundary

    sk = _anchor_key(start_anchor)
    for item in build_global_section_graph(pages, custom_headings=custom_headings):
        if _anchor_key(item) <= sk: continue
        if item.get("generic_boundary") and not item.get("hard_boundary"): continue
        if _boundary_is_child(requested_label, item, pages): continue
        return item
    return None

def _leadership_end_anchor(pages, start_anchor, boundary_anchor, label):
    role_terms = {"Chairman Message": ["chairman", "chairperson"], "CEO Message": ["chief executive officer", "ceo"], "Managing Director Message": ["managing director"]}.get(label, [])
    if not role_terms: return boundary_anchor

    stop_idx = boundary_anchor["index"] if boundary_anchor else min(len(pages), start_anchor["index"] + 12)
    last_signature_idx = None
    for idx in range(start_anchor["index"], min(stop_idx, start_anchor["index"] + 12)):
        lines = [_norm_line(x.get("text", "")) for x in _page_lines(pages[idx]) if _norm_line(x.get("text", ""))]
        if not lines: continue
        tail = " ".join(lines[-30:]).lower().replace("’", "'")
        if any(term in tail for term in role_terms): last_signature_idx = idx
    if last_signature_idx is None: return boundary_anchor

    ni = last_signature_idx + 1
    if boundary_anchor and ni > boundary_anchor["index"]: return boundary_anchor
    if ni >= len(pages): return None
    return {"label": "Leadership signature end", "index": ni, "pdf_page": pages[ni].get("page"), "line_order": 0, "printed_page": None, "score": 100}

def _selected_line_text(page, start_order=None, end_order=None, clean_page=None, x_region=None):
    lines = _page_lines(page)
    if not lines: return page.get("text", "").strip()
    clean_allowed = Counter(_norm_line(x) for x in clean_page.get("text", "").splitlines() if _norm_line(x)) if clean_page is not None else None
    chosen = []
    for line in lines:
        order = line.get("order", 0)
        if start_order is not None and order < start_order: continue
        if end_order is not None and order >= end_order: continue
        bb = line.get("bbox")
        if x_region is not None and bb:
            xc = (bb[0] + bb[2]) / 2
            if not (x_region[0] <= xc <= x_region[1]): continue
        text = _norm_line(line.get("text", ""))
        if not text: continue
        if clean_allowed is not None:
            if clean_allowed[text] <= 0: continue
            clean_allowed[text] -= 1
        chosen.append(text)
    return "\n".join(chosen).strip()

def _logical_side(page, anchor):
    if not anchor or not anchor.get("bbox"): return None
    width = float(page.get("width") or 0); height = float(page.get("height") or 0)
    if not width or not height or width / max(height, 1) < 1.25: return None
    folios = _edge_folios(page)
    if len({f["number"] for f in folios}) < 2: return None
    xc = (anchor["bbox"][0] + anchor["bbox"][2]) / 2
    return "left" if xc < width / 2 else "right"

def _half_region(page, side):
    width = float(page.get("width") or 0)
    if not width or side not in {"left", "right"}: return None
    pad = 0.015 * width; mid = width / 2
    return (0, mid + pad) if side == "left" else (mid - pad, width)

def _payload_from_anchors(raw_pages, clean_pages_, start_anchor, boundary_anchor=None):
    start_idx = start_anchor["index"]
    boundary_idx = boundary_anchor["index"] if boundary_anchor else len(raw_pages)
    raw_parts, clean_parts, used_pages = [], [], []

    start_side = _logical_side(raw_pages[start_idx], start_anchor) if start_idx < len(raw_pages) else None
    boundary_side = _logical_side(raw_pages[boundary_idx], boundary_anchor) if boundary_anchor and boundary_idx < len(raw_pages) else None

    include_boundary_page = False; boundary_region = None
    if boundary_anchor and boundary_idx < len(raw_pages):
        ps = start_anchor.get("printed_page"); bp = boundary_anchor.get("printed_page")
        folios = {f["number"] for f in _edge_folios(raw_pages[boundary_idx])}
        if boundary_side == "right" and bp is not None and (bp - 1) in folios:
            include_boundary_page = True; boundary_region = _half_region(raw_pages[boundary_idx], "left")
        elif boundary_side == "left": include_boundary_page = False
        elif ps is not None and bp is not None and bp > ps and (bp - 1) in folios and bp in folios: include_boundary_page = True

    same_page_boundary = bool(
        boundary_anchor and boundary_idx == start_idx
        and boundary_anchor.get("line_order", 0) > start_anchor.get("line_order", 0)
        and not include_boundary_page
    )
    if same_page_boundary:
        last_idx = start_idx
    else:
        last_idx = boundary_idx if (boundary_anchor and include_boundary_page) else (boundary_idx - 1 if boundary_anchor else len(raw_pages) - 1)
    for idx in range(start_idx, last_idx + 1):
        if idx >= len(raw_pages): break
        start_order = start_anchor.get("line_order") if idx == start_idx else None
        x_region = None
        if idx == start_idx and start_anchor.get("detection_source") in {"native", "ocr-probe"}: start_order = 0
        if idx == start_idx and start_side == "right": x_region = _half_region(raw_pages[idx], "right"); start_order = 0

        end_order = None
        if same_page_boundary and idx == start_idx:
            end_order = boundary_anchor.get("line_order")
        elif boundary_anchor and include_boundary_page and idx == boundary_idx:
            if boundary_region is not None: x_region = boundary_region; end_order = None
            else: end_order = boundary_anchor.get("line_order")

        r = _selected_line_text(raw_pages[idx], start_order=start_order, end_order=end_order, x_region=x_region)
        c = _selected_line_text(raw_pages[idx], start_order=start_order, end_order=end_order, clean_page=clean_pages_[idx], x_region=x_region)
        if r.strip():
            raw_parts.append(r.strip()); clean_parts.append(c.strip() if c.strip() else r.strip()); used_pages.append(raw_pages[idx].get("page"))

    if not used_pages: return None
    payload = {
        "start_page": min(p for p in used_pages if p is not None) if any(p is not None for p in used_pages) else None,
        "end_page": max(p for p in used_pages if p is not None) if any(p is not None for p in used_pages) else None,
        "text": "\n\n".join(clean_parts).strip(),
        "raw_text": "\n\n".join(raw_parts).strip(),
        "detection_confidence": "high" if start_anchor.get("score", 0) >= 64 else "medium",
    }
    detected_heading = _norm_line(start_anchor.get("matched_text", ""))
    if detected_heading:
        early = _compact(" ".join(payload["text"].splitlines()[:10]))
        alias_c = _compact(start_anchor.get("matched_alias", detected_heading))
        if alias_c and alias_c not in early: payload["text"] = detected_heading + "\n" + payload["text"]; payload["raw_text"] = detected_heading + "\n" + payload["raw_text"]
    ps = start_anchor.get("printed_page")
    if ps is not None: payload["printed_start_page"] = ps

    pe = None
    if boundary_anchor and ps is not None and boundary_anchor.get("printed_page") is not None and boundary_anchor["printed_page"] > ps: pe = boundary_anchor["printed_page"] - 1
    elif ps is not None:
        vals = []
        used_set = set(used_pages)
        for page in raw_pages:
            if page.get("page") not in used_set: continue
            vals.extend(f["number"] for f in _edge_folios(page) if ps <= f["number"] <= ps + 500)
        if vals: pe = max(vals)
    if pe is not None: payload["printed_end_page"] = pe
    return payload

def _graph_start_for_label(pages, label, custom_headings=None):
    target = _canonical_request_label(label)
    candidates = [n for n in build_global_section_graph(pages, custom_headings=custom_headings) if _canonical_request_label(n.get("label")) == target]
    if not candidates: return None
    candidates.sort(key=lambda n: (-n.get("score", 0), _anchor_key(n)))
    return dict(candidates[0])

def extract_preset(raw_pages, clean_pages_, label):
    cfg = PRESETS[label]
    responsibility_labels = {"Business Responsibility Report (BRR)", "Business Responsibility & Sustainability Report (BRSR)", "ESG Report", "Sustainability Report"}
    separate_status = _separate_enclosure_status(raw_pages, cfg["aliases"]) if label in responsibility_labels else None

    start = None
    if separate_status == "web-only": return None

    if label in {"Business Responsibility Report (BRR)", "Business Responsibility & Sustainability Report (BRSR)"}:
        start = _native_visual_wrapper_anchor(clean_pages_, cfg["aliases"], kind="responsibility")
        if not start and raw_pages is not clean_pages_:
            start = _native_visual_wrapper_anchor(raw_pages, cfg["aliases"], kind="responsibility")

    if separate_status == "separate" and not start:
        start = _find_separate_appendix_anchor(clean_pages_, cfg["aliases"])
        if not start: return None
    elif not start:
        start = _detect_section_anchor(clean_pages_, cfg["aliases"])
        if not start: start = _native_exact_heading_anchor(clean_pages_, cfg["aliases"])
        if not start: start = _graph_start_for_label(clean_pages_, label)
        if not start and raw_pages is not clean_pages_:
            start = _detect_section_anchor(raw_pages, cfg["aliases"])
            if not start: start = _native_exact_heading_anchor(raw_pages, cfg["aliases"])
            if not start: start = _graph_start_for_label(raw_pages, label)
    if not start: return None

    boundary = _next_boundary_anchor(clean_pages_, start, requested_label=label)
    if label in {"Chairman Message", "CEO Message", "Managing Director Message"}:
        preserve_logical_boundary = bool(boundary and _is_two_up_document(clean_pages_) and start.get("printed_page") is not None and boundary.get("printed_page") is not None and boundary.get("printed_page") > start.get("printed_page"))
        if not preserve_logical_boundary: boundary = _leadership_end_anchor(clean_pages_, start, boundary, label)
    return _payload_from_anchors(raw_pages, clean_pages_, start, boundary)

def extract_custom(raw_pages, clean_pages_, heading, custom_headings=None):
    heading = _norm_line(heading)
    if not heading: return None
    custom_headings = [_norm_line(x) for x in (custom_headings or [heading]) if _norm_line(x)]
    if heading not in custom_headings: custom_headings.append(heading)

    alias_map = {
        _compact("Directors' Report"): ["directors' report", "directors’ report", "director's report", "directors' report to the shareholders"],
        _compact("Board's Report"): ["board's report", "boards' report", "boards’ report"],
        _compact("Corporate Governance Report"): ["corporate governance report", "report on corporate governance"],
        _compact("Corporate social responsibility"): ["corporate social responsibility", "corporate social responsibility report", "csr report"],
    }
    aliases = alias_map.get(_compact(heading), [heading])
    requested = "Board's Report" if _compact(heading) == _compact("Board's Report") else ("Directors' Report" if _compact(heading) == _compact("Directors' Report") else heading)

    start = None
    if requested in {"Board's Report", "Directors' Report"}:
        start = _native_visual_wrapper_anchor(clean_pages_, aliases, kind="board")
        if not start and raw_pages is not clean_pages_:
            start = _native_visual_wrapper_anchor(raw_pages, aliases, kind="board")
    if not start: start = _detect_section_anchor(clean_pages_, aliases)
    if not start: start = _native_exact_heading_anchor(clean_pages_, aliases)
    if not start: start = _graph_start_for_label(clean_pages_, requested, custom_headings=custom_headings)
    if not start and raw_pages is not clean_pages_:
        start = _detect_section_anchor(raw_pages, aliases)
        if not start: start = _native_exact_heading_anchor(raw_pages, aliases)
        if not start: start = _graph_start_for_label(raw_pages, requested, custom_headings=custom_headings)
    if not start: return None

    boundary = _next_boundary_anchor(clean_pages_, start, requested_label=requested, custom_headings=custom_headings)
    return _payload_from_anchors(raw_pages, clean_pages_, start, boundary)

def combine_sections(raw_pages, clean_pages_, sections, labels, combined_label):
    present = [(label, sections[label]) for label in labels if label in sections]
    if not present: return None
    clean_parts, raw_parts, page_ranges, printed_ranges = [], [], [], []
    for label, sec in present:
        if sec.get("text", "").strip(): clean_parts.append(sec["text"].strip()); raw_parts.append(sec.get("raw_text", sec["text"]).strip())
        if sec.get("start_page") is not None: page_ranges.append({"label": label, "start_page": sec.get("start_page"), "end_page": sec.get("end_page")})
        if sec.get("printed_start_page") is not None: printed_ranges.append({"label": label, "start_page": sec.get("printed_start_page"), "end_page": sec.get("printed_end_page", sec.get("printed_start_page"))})
    payload = {
        "start_page": min((r["start_page"] for r in page_ranges), default=None), "end_page": max((r["end_page"] for r in page_ranges), default=None),
        "page_ranges": page_ranges, "text": "\n\n".join(clean_parts).strip(), "raw_text": "\n\n".join(raw_parts).strip(), "combined_from": [label for label, _ in present],
    }
    if printed_ranges:
        payload["printed_page_ranges"] = printed_ranges
        payload["printed_start_page"] = min(r["start_page"] for r in printed_ranges); payload["printed_end_page"] = max(r["end_page"] for r in printed_ranges)
    return payload

def _preserved_page_text(page):
    native = page.get("native_text")
    if native is not None and len(str(native).strip()) >= 20 and _text_quality(native) >= 0.45: return str(native).strip()
    return str(page.get("text", "")).strip()

def raw_full_text(pages):
    if len(pages) == 1 and pages[0].get("page") is None: return pages[0]["text"].strip()
    return "\n\n".join(f"===== PAGE {p['page']} =====\n\n{_preserved_page_text(p)}" for p in pages).strip()

def search_pages(pages, query, limit=50):
    query = query.strip()
    if not query: return []
    pat = re.compile(re.escape(query), re.I)
    hits = []
    for p in pages:
        for m in pat.finditer(p["text"]):
            a = max(0, m.start() - 90); b = min(len(p["text"]), m.end() + 150)
            hits.append({"page": p.get("page"), "snippet": re.sub(r"\s+", " ", p["text"][a:b]).strip()})
            if len(hits) >= limit: return hits
    return hits

def base_stem(filename):
    stem = filename.rsplit(".", 1)[0]
    stem = re.sub(r"(?:_Annual_Report)?_Full_Text$", "", stem, flags=re.I)
    return stem

def _normalise_entity_name(name):
    name = re.sub(r"\s+", " ", str(name or "")).strip(" ,.-|:")
    return re.sub(r"(?i)^(?:name\s+of\s+(?:the\s+)?(?:company|listed\s+entity)|listed\s+entity|company\s+name|for\s+and\s+on\s+behalf\s+of|for)\s*[:\-–—]*\s*", "", name).strip(" ,.-|:")

def _entity_key(name):
    return _compact(re.sub(r"(?i)\b(?:private\s+limited|limited|ltd\.?)\b", "", name))

def _role_context_penalty(context):
    role_terms = ("registrar", "share transfer agent", "transfer agent", "registrar and transfer agent", "statutory auditor", "secretarial auditor", "internal auditor", "cost auditor", "auditor", "chartered accountants", "banker", "depository", "trustee", "assurance provider", "verification provider", "scrutinizer", "legal counsel", "subsidiary", "associate company", "joint venture", "step-down subsidiary", "stock exchange", "exchange plaza", "nominee director")
    low = context.lower().replace("’", "'")
    return min(260.0, 55.0 * sum(1 for t in role_terms if t in low))

def _entity_positive_context(context):
    low = context.lower().replace("’", "'")
    positive = (("integrated annual report", 26), ("annual report", 20), ("annual general meeting", 24), ("registered office", 18), ("corporate identity number", 30), ("cin", 18), ("name of the company", 45), ("name of the listed entity", 55), ("listed entity", 20), ("the company", 8))
    return min(100.0, sum(pts for term, pts in positive if term in low))

def _entity_candidates_from_line(text):
    text = _norm_line(text)
    if not text: return []
    explicit = re.search(r"(?i)(?:name\s+of\s+(?:the\s+)?(?:company|listed\s+entity)|company\s+name)\s*[:\-–—]*\s*([A-Z][A-Za-z0-9&.,'()\-/ ]{2,120}?(?:Private\s+Limited|Limited|Ltd\.?))\b", text)
    out = []
    if explicit: out.append((_normalise_entity_name(explicit.group(1)), "explicit-name-field"))
    entity_re = re.compile(r"([A-Z][A-Za-z0-9&.,'()\-/ ]{2,110}?(?:Private\s+Limited|Limited|Ltd\.?))\b")
    for m in entity_re.finditer(text):
        name = _normalise_entity_name(m.group(1))
        if name: out.append((name, "legal-name"))
    seen = set(); result = []
    for name, source in out:
        key = _entity_key(name)
        if key and key not in seen: seen.add(key); result.append((name, source))
    return result

def score_company_metadata(pages):
    stats = {}
    for pi, page in enumerate(pages[:60]):
        page_no = page.get("page") or (pi + 1); height = float(page.get("height") or 800.0)
        lines = page.get("native_lines") or _page_lines(page)
        if not lines: lines = [{"text": x, "bbox": None, "size": None, "bold": False} for x in page.get("text", "").splitlines()]
        line_texts = [_norm_line(x.get("text", "")) for x in lines]
        page_sizes = [x.get("size") for x in lines if isinstance(x.get("size"), (int, float))]
        page_max = max(page_sizes) if page_sizes else 0.0

        for li, line in enumerate(lines):
            text = _norm_line(line.get("text", ""))
            if not text: continue
            context = " ".join(line_texts[max(0, li - 3): min(len(line_texts), li + 4)])
            for raw, source in _entity_candidates_from_line(text):
                key = _entity_key(raw)
                if len(key) < 4: continue
                rec = stats.setdefault(key, {"name_scores": Counter(), "score": 0.0, "pages": set(), "margin_pages": set(), "occurrences": 0, "formal_hits": 0})
                rec["name_scores"][raw] += 1; rec["occurrences"] += 1; rec["pages"].add(page_no); rec["score"] += 5
                if source == "explicit-name-field": rec["score"] += 75; rec["formal_hits"] += 1
                if page_no <= 2: rec["score"] += 18
                elif page_no <= 5: rec["score"] += 10
                elif page_no <= 10: rec["score"] += 4

                size = line.get("size")
                if isinstance(size, (int, float)):
                    if page_max and size >= max(20.0, 0.92 * page_max): rec["score"] += 42
                    elif size >= 16: rec["score"] += 24
                    elif size >= 12: rec["score"] += 8
                if line.get("bold"): rec["score"] += 4

                bb = line.get("bbox")
                if bb and height and page_no <= 30:
                    near_margin = bb[1] <= 0.08 * height or bb[3] >= 0.92 * height
                    if near_margin: rec["margin_pages"].add(page_no); rec["score"] += 7

                rec["score"] += _entity_positive_context(context); rec["score"] -= _role_context_penalty(context)
                if re.search(r"(?i)\bfor\s+" + re.escape(raw), context): rec["score"] += 16
    if not stats: return None
    for rec in stats.values():
        rec["score"] += min(150, 6 * len(rec["pages"])) + min(100, 5 * len(rec["margin_pages"])) + min(60, 12 * rec["formal_hits"])
    keys = list(stats)
    for k in keys:
        for other in keys:
            if k == other: continue
            if len(k) < len(other) and (other.endswith(k) or k in other) and len(other) - len(k) >= 4:
                if len(stats[other]["pages"]) >= max(2, len(stats[k]["pages"]) // 2):
                    stats[k]["score"] -= 40; break

    winner_key, winner = max(stats.items(), key=lambda kv: (kv[1]["score"], len(kv[1]["margin_pages"]), len(kv[1]["pages"]), len(kv[0])))
    names = list(winner["name_scores"])
    names.sort(key=lambda n: (winner["name_scores"][n], len(n)), reverse=True)
    return names[0] if names else None

def _canonical_year_range(a, b):
    a = int(a); btxt = str(b)
    bend = int(btxt[-2:]) if len(btxt) <= 2 else int(btxt) % 100
    return f"{a:04d}-{bend:02d}"

def _infer_financial_year(filename, pages):
    early, priority = [], []
    for i, page in enumerate(pages[:60]):
        native = " ".join(x.get("text", "") for x in (page.get("native_lines") or []))
        text = (page.get("text", "") + "\n" + native).strip()
        if i < 10: early.append(text)
        if i < 10 or any(k in text.lower() for k in ("notice", "auditor's report", "auditors' report", "annual report")): priority.append(text)
    early_text, priority_text = "\n".join(early), "\n".join(priority)

    range_patterns = [r"(?i)(?:integrated\s+)?annual\s+report(?:\s*&\s*accounts)?(?:\s+for\s+the\s+year)?\s*[:\-]?\s*(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})", r"(?i)\bFY\s*[:\-]?\s*(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})", r"(?i)financial\s+year(?:\s+reported|\s+for\s+which\s+reporting\s+is\s+being\s+done)?\s*[:\-]?\s*(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})"]
    for corpus in (early_text, priority_text):
        for pat in range_patterns:
            m = re.search(pat, corpus, re.S)
            if m: return _canonical_year_range(m.group(1), m.group(2)), "high"

    march_patterns = [r"(?i)for\s+the\s+(?:financial\s+)?year\s+ended\s+(?:on\s+)?31(?:st)?\s+march[,]?\s*(20\d{2})", r"(?i)year\s+ended\s+(?:on\s+)?(?:march\s+31|31(?:st)?\s+march)[, ]+\s*(20\d{2})"]
    for pat in march_patterns:
        m = re.search(pat, priority_text)
        if m: end_year = int(m.group(1)); return f"{end_year - 1:04d}-{end_year % 100:02d}", "high"

    pairs = re.findall(r"\b(20\d{2})\s*[-–—_/]\s*(\d{2,4})\b", early_text)
    if pairs: ranges = [_canonical_year_range(a, b) for a, b in pairs]; return Counter(ranges).most_common(1)[0][0], "medium"
    m = re.search(r"\b(20\d{2})\s*[-–—_/]\s*(\d{2,4})\b", filename)
    if m: return _canonical_year_range(m.group(1), m.group(2)), "medium"
    m = re.search(r"\b(20\d{2})\b", filename)
    if m: return m.group(1), "low"
    return "Year not detected", "none"

def infer_metadata(filename, pages):
    stem = base_stem(filename)
    year, year_confidence = _infer_financial_year(filename, pages)
    company = score_company_metadata(pages)
    if not company:
        company = re.sub(r"\s*\(\d+\)\s*$", "", stem)
        company = re.sub(r"[_-]+", " ", company)
        company = re.sub(r"\b(?:Integrated\s+)?Annual\s+Report\b", " ", company, flags=re.I)
        company = re.sub(r"\b20\d{2}(?:\s*[-–—_/]?\s*\d{2,4})?\b", " ", company)
    company = _normalise_entity_name(company)
    company = re.sub(r"(?i)\s+(?:Private\s+Limited|Limited|Ltd\.?)$", "", company).strip(" ,.-")
    return {"company": company or stem.replace("_", " ").strip(), "year": year, "year_confidence": year_confidence, "source_file": filename}

def _docx_bytes(title, text):
    d = Document(); d.add_heading(title, 0)
    for block in text.split("\n\n"):
        if block.strip(): d.add_paragraph(block.strip())
    b = io.BytesIO(); d.save(b)
    return b.getvalue()

def _csv_bytes(rows):
    s = io.StringIO(); w = csv.DictWriter(s, fieldnames=rows[0].keys()); w.writeheader(); w.writerows(rows)
    return s.getvalue().encode("utf-8-sig")

def _pdf_bytes(title, text):
    buf = io.BytesIO(); writer = fitz.DocumentWriter(buf); mediabox = fitz.paper_rect("a4"); rect = fitz.Rect(42, 42, mediabox.width - 42, mediabox.height - 42)
    body = html.escape(text).replace("\n", "<br>")
    story = fitz.Story(html=f"<h2>{html.escape(title)}</h2><div>{body}</div>", user_css="body {font-family: sans-serif; font-size: 9pt; line-height: 1.25;} h2 {font-size: 14pt;}")
    more = True
    while more: dev = writer.begin_page(mediabox); more, _ = story.place(rect); story.draw(dev); writer.end_page()
    writer.close()
    return buf.getvalue()

def safe_label(label):
    return {"Full Report": "Annual_Report_Full_Text", "Leadership Messages (Combined)": "Chairman_CEO_MD_Messages", "Chairman Message": "Chairman_Message", "CEO Message": "CEO_Message", "Managing Director Message": "Managing_Director_Message", "Management Discussion & Analysis": "MDA", "Business Responsibility & Sustainability Report (BRSR)": "BRSR", "Business Responsibility Report (BRR)": "BRR", "ESG Report": "ESG_Report", "Sustainability Report": "Sustainability_Report"}.get(label, re.sub(r"[^A-Za-z0-9_-]+", "_", label).strip("_"))

_safe_label = safe_label

def format_file(stem, label, payload, fmt, raw_pages, source_bytes=None, source_is_pdf=False, all_sections=None):
    safe = safe_label(label); name = f"{stem}_{safe}"; text = payload["text"]
    if fmt == "TXT": return name + ".txt", text.encode("utf-8")
    if fmt == "MD": return name + ".md", (f"# {label}\n\n{text}\n").encode("utf-8")
    if fmt == "DOCX": return name + ".docx", _docx_bytes(label, text)
    if fmt == "PDF":
        if label == "Full Report" and source_is_pdf and source_bytes is not None: return name + ".pdf", source_bytes
        return name + ".pdf", _pdf_bytes(label, text)
    if fmt == "JSON":
        if label == "Full Report": obj = {"file": stem, "page_count": len(raw_pages) if raw_pages and raw_pages[0].get("page") is not None else None, "pages": raw_pages, "text": text, "sections": {k: {"start_page": v.get("start_page"), "end_page": v.get("end_page")} for k, v in (all_sections or {}).items()}}
        else: obj = payload
        return name + ".json", json.dumps(obj, ensure_ascii=False, indent=2).encode("utf-8")
    if fmt == "CSV":
        if label == "Full Report": rows = [{"page": p.get("page"), "text": p["text"]} for p in raw_pages]
        else: rows = [{"section": label, "start_page": payload.get("start_page"), "end_page": payload.get("end_page"), "text": text}]
        return name + ".csv", _csv_bytes(rows)
    raise ValueError(fmt)

def make_zip(files):
    b = io.BytesIO()
    with zipfile.ZipFile(b, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items(): z.writestr(name, data)
    return b.getvalue()

def manifest_csv(results):
    s = io.StringIO(); fields = ["company", "year", "source_file", "section", "start_page", "end_page", "formats"]; w = csv.DictWriter(s, fieldnames=fields); w.writeheader()
    for r in results:
        w.writerow({"company": r["meta"]["company"], "year": r["meta"]["year"], "source_file": r["name"], "section": "Full Report", "start_page": 1 if len(r["raw_pages"]) > 1 else "", "end_page": len(r["raw_pages"]) if len(r["raw_pages"]) > 1 else "", "formats": ", ".join(r["formats"])})
        for label, sec in r["sections"].items(): w.writerow({"company": r["meta"]["company"], "year": r["meta"]["year"], "source_file": r["name"], "section": label, "start_page": sec.get("start_page", ""), "end_page": sec.get("end_page", ""), "formats": ", ".join(r["formats"])})
    return s.getvalue().encode("utf-8-sig")

def clean_background_pdf(pdf_bytes, threshold=215):
    src = fitz.open(stream=pdf_bytes, filetype="pdf"); out = fitz.open(); lut = [x if x < threshold else 255 for x in range(256)]
    for p in src:
        pix = p.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), colorspace=fitz.csGRAY, alpha=False); img = Image.frombytes("L", (pix.width, pix.height), pix.samples).point(lut); b = io.BytesIO(); img.save(b, "PNG"); q = out.new_page(width=p.rect.width, height=p.rect.height); q.insert_image(q.rect, stream=b.getvalue())
    return out.tobytes(deflate=True)