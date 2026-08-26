import csv
import html
import io
import json
import re
import zipfile
from collections import Counter
from copy import deepcopy

import fitz
import pytesseract
from PIL import Image
from docx import Document

PRESETS = {
    "Chairman Message": {
        "aliases": [
            "message from the chairman", "chairman's message", "chairperson's message",
            "message from the chairperson", "letter from the chairman",
            "chairman's statement", "chairman statement", "statement from the chairman",
        ],
        "end": [
            "message from the managing director", "managing director's message",
            "message from the ceo", "ceo message", "our leadership",
            "board of directors", "team of executives", "contents", "index",
            "about the company", "about adani enterprises", "company overview", "company profile",
        ],
    },
    "CEO Message": {
        "aliases": ["message from the ceo", "ceo message", "ceo's message", "letter from the ceo"],
        "end": [
            "message from the managing director", "managing director's message",
            "our leadership", "board of directors", "team of executives", "contents", "index",
            "about the company", "company overview", "company profile", "business portfolio",
        ],
    },
    "Managing Director Message": {
        "aliases": [
            "message from the managing director", "managing director's message",
            "md message", "letter from the managing director",
        ],
        "end": [
            "key performance indicators", "our growth engines", "about adani enterprises",
            "about the company", "company overview", "company profile", "business portfolio",
            "board of directors", "contents", "index",
        ],
    },
    "Management Discussion & Analysis": {
        "aliases": [
            "management discussion & analysis report",
            "management discussion and analysis report",
            "management discussion & analysis",
            "management discussion and analysis",
            "md&a",
        ],
        "end": [
            "annexures i to vi to the directors' report",
            "annexures i to vi to the directors’ report",
            "annexure i to vi to the director's report",
            "annexures to the directors' report",
            "annexures to the directors’ report",
            "corporate governance report", "corporate governance",
        ],
    },
    "Business Responsibility & Sustainability Report (BRSR)": {
        "aliases": [
            "business responsibility & sustainability report",
            "business responsibility and sustainability report", "brsr report", "brsr",
        ],
        "end": [
            "standalone financial statements", "independent auditor's report",
            "independent auditors' report", "report on the audit of the standalone",
        ],
    },
    "Business Responsibility Report (BRR)": {
        "aliases": ["business responsibility report", "business responsibility (br) report", "brr report"],
        "end": [
            "standalone financial statements", "independent auditor's report",
            "independent auditors' report", "report on the audit of the standalone",
        ],
    },
    "ESG Report": {
        "aliases": ["esg report", "environmental social and governance report"],
        "end": [
            "standalone financial statements", "independent auditor's report",
            "independent auditors' report", "report on the audit of the standalone",
        ],
    },
    "Sustainability Report": {
        "aliases": ["sustainability report"],
        "end": [
            "standalone financial statements", "independent auditor's report",
            "independent auditors' report", "report on the audit of the standalone",
        ],
    },
}

FORMAT_EXT = {"TXT": "txt", "JSON": "json", "PDF": "pdf", "DOCX": "docx", "CSV": "csv", "MD": "md"}


def _compact(s):
    return re.sub(r"[^a-z0-9]+", "", s.lower())


def _norm_line(s):
    return re.sub(r"\s+", " ", s).strip()


def _ocr(page):
    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), colorspace=fitz.csGRAY, alpha=False)
    img = Image.frombytes("L", (pix.width, pix.height), pix.samples)
    return pytesseract.image_to_string(img, config="--psm 6")


def extract_source(name, data):
    ext = name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            method = "native"
            if len(text.strip()) < 40:
                text = _ocr(page)
                method = "ocr"
            pages.append({"page": i + 1, "text": text, "method": method})
        return pages

    if ext in ("txt", "md"):
        text = data.decode("utf-8", errors="ignore")
        parts = re.split(r"(?m)^\s*=====\s*PAGE\s+(\d+)\s*=====\s*$", text)
        if len(parts) > 2:
            return [
                {"page": int(parts[i]), "text": parts[i + 1], "method": "text"}
                for i in range(1, len(parts), 2)
            ]
        return [{"page": None, "text": text, "method": "text"}]

    if ext == "docx":
        doc = Document(io.BytesIO(data))
        return [{"page": None, "text": "\n".join(p.text for p in doc.paragraphs), "method": "docx"}]

    raise ValueError(f"Unsupported file type: {ext}")


def clean_pages(pages):
    """Research-safe cleanup.

    Removes only recurring TEXTUAL headers/footers when they occur at page edges.
    Numeric/table values are preserved. Raw pages are never modified.
    """
    out = deepcopy(pages)
    if len(out) <= 1:
        out[0]["text"] = "\n".join(
            _norm_line(x) for x in out[0]["text"].splitlines() if _norm_line(x)
        )
        return out

    edge = Counter()
    n = len(out)

    for p in out:
        lines = [_norm_line(x) for x in p["text"].splitlines() if _norm_line(x)]
        edge_lines = lines[:7] + lines[-7:]
        for x in set(edge_lines):
            words = re.findall(r"[A-Za-z]+", x)
            # Never classify pure numbers / short table values as boilerplate.
            if 4 < len(x) <= 140 and len(words) >= 2:
                edge[x] += 1

    repeated = {x for x, c in edge.items() if c >= max(3, int(0.20 * n))}

    for p in out:
        original = [_norm_line(x) for x in p["text"].splitlines() if _norm_line(x)]
        cleaned = []
        last = len(original) - 1
        for i, x in enumerate(original):
            at_edge = i < 7 or i > last - 7
            if at_edge and x in repeated:
                continue
            cleaned.append(x)
        p["text"] = "\n".join(cleaned)

    return out


COMMON_SECTION_HEADINGS = [
    "directors' report", "board's report",
    "management discussion and analysis", "management discussion & analysis",
    "annexures", "corporate governance report",
    "business responsibility report", "business responsibility & sustainability report",
    "standalone financial statements", "consolidated financial statements",
    "independent auditor's report", "notice", "gri index",
]

COMMON_SECTION_COMPACTS = [_compact(h) for h in COMMON_SECTION_HEADINGS]

def _toc_like_page(text):
    low = text.lower()
    if "contents" in low:
        return True
    compact_text = _compact(text)
    hits = sum(h in compact_text for h in COMMON_SECTION_COMPACTS)
    return hits >= 5

def _heading_variants(heading):
    if isinstance(heading, (list, tuple, set)):
        raw = list(heading)
    else:
        raw = [heading]

    out = []
    for h in raw:
        if not h:
            continue
        h = _norm_line(str(h))
        variants = [h]
        variants.append(re.sub(r"\s*\([^)]*\)\s*", " ", h).strip())
        if "&" in h:
            variants.append(h.replace("&", "and"))
        if re.search(r"\band\b", h, re.I):
            variants.append(re.sub(r"\band\b", "&", h, flags=re.I))
        if re.search(r"\breport\b$", h, re.I):
            variants.append(re.sub(r"\breport\b$", "", h, flags=re.I).strip())
        out.extend(v for v in variants if v)

    # Preserve order, remove duplicate compact forms.
    seen = set()
    unique = []
    for h in out:
        c = _compact(h)
        if c and c not in seen:
            seen.add(c)
            unique.append(h)
    return unique


def _toc_printed_page(pages, heading):
    """Read a printed report page number from a Contents/Index page.

    Supports both common layouts:
      Management Discussion and Analysis Report .... 17
      17 .... Management Discussion and Analysis Report
    """
    variants = _heading_variants(heading)

    for p in pages[:25]:
        text = p["text"]
        low = text.lower()
        if "contents" not in low and "\nindex\n" not in f"\n{low}\n" and " index " not in f" {low} ":
            continue

        flat = re.sub(r"\s+", " ", text.replace("…", "."))
        for variant in variants:
            words = re.findall(r"[A-Za-z0-9]+", variant)
            if not words:
                continue

            # Allow OCR/layout punctuation between heading words.
            phrase = r"[\s&'’()./\-–—:]*".join(re.escape(w) for w in words)
            sep = r"[\s.\-–—:·•]*"

            # Heading first, page number after it (the common annual-report TOC form).
            m = re.search(rf"{phrase}{sep}\b(0*\d{{1,4}})\b", flat, re.I)
            if m:
                return int(m.group(1))

    return None


def _edge_printed_page_numbers(text):
    """Return likely printed page numbers from page header/footer text only."""
    lines = [_norm_line(x) for x in text.splitlines() if _norm_line(x)]
    edge = lines[:24] + lines[-24:]
    nums = set()

    for line in edge:
        # Clean standalone printed page number: 17, 02, etc.
        m = re.fullmatch(r"0*(\d{1,4})", line)
        if m:
            nums.add(int(m.group(1)))
            continue

        # Footer/header forms such as:
        # "123rd Annual Report & Accounts 2019-20 03"
        # "03 Century Textiles and Industries Limited"
        if len(line) <= 140 and (
            "annual report" in line.lower()
            or "century textiles" in line.lower()
            or "company overview" in line.lower()
            or "statutory reports" in line.lower()
            or "financial statements" in line.lower()
        ):
            m = re.match(r"^0*(\d{1,4})\b", line)
            if m:
                nums.add(int(m.group(1)))
            m = re.search(r"\b0*(\d{1,4})$", line)
            if m:
                nums.add(int(m.group(1)))
    return nums


def _contains_printed_page_number(text, number):
    if number is None:
        return False
    return int(number) in _edge_printed_page_numbers(text)


def _physical_index_for_printed_page(pages, printed_page, start_after=-1):
    if printed_page is None:
        return None
    for i in range(start_after + 1, len(pages)):
        if _toc_like_page(pages[i]["text"]):
            continue
        if _contains_printed_page_number(pages[i]["text"], printed_page):
            return i
    return None


def _heading_score(text, heading, max_lines=180):
    """Score whether a phrase is functioning as a page heading, not merely a body reference."""
    target = _compact(heading)
    lines = [_norm_line(x) for x in text.splitlines()[:max_lines] if _norm_line(x)]
    compact = [_compact(x) for x in lines]
    best = 0

    for i in range(len(compact)):
        # Exact one-line heading.
        if compact[i] == target:
            score = 10
            if lines[i].isupper():
                score += 15
            if i < 20:
                score += 8
            best = max(best, score)

        # PDF extraction often splits a heading over 2-3 lines.
        for width in (2, 3):
            if i + width <= len(compact) and "".join(compact[i:i+width]) == target:
                score = 9
                joined_original = " ".join(lines[i:i+width])
                if joined_original.isupper():
                    score += 15
                if i < 20:
                    score += 8
                best = max(best, score)

    return best

BOUNDARY_SECTIONS = [
    ("Directors' Report", ["directors' report", "directors’ report", "director's report", "board's report", "board’s report"]),
    ("Management Discussion & Analysis", PRESETS["Management Discussion & Analysis"]["aliases"]),
    ("Annexures", [
        "annexures i to vi to the directors' report", "annexures i to vi to the directors’ report",
        "annexure i to vi to the director's report", "annexures to the directors' report",
        "annexures to the directors’ report", "annexures i to vi", "annexures",
    ]),
    ("Corporate Governance Report", ["corporate governance report", "report on corporate governance"]),
    ("BRSR", PRESETS["Business Responsibility & Sustainability Report (BRSR)"]["aliases"]),
    ("BRR", PRESETS["Business Responsibility Report (BRR)"]["aliases"]),
    ("ESG / Sustainability Report", PRESETS["ESG Report"]["aliases"] + PRESETS["Sustainability Report"]["aliases"]),
    ("Standalone Financial Statements", ["standalone financial statements", "financial statements - standalone"]),
    ("Consolidated Financial Statements", ["consolidated financial statements", "financial statements - consolidated"]),
    ("Independent Auditor's Report", [
        "independent auditor's report", "independent auditors' report", "independent auditor’s report",
        "independent auditors’ report", "report on the audit of the standalone financial statements",
    ]),
    ("Notice", ["notice", "notice of annual general meeting"]),
    ("GRI Index", ["gri index", "gri content index"]),
]

_SECTION_MAP_CACHE = {}


def _doc_key(pages):
    """Cheap cache key; avoids rebuilding the same section map for every checkbox."""
    if not pages:
        return (0,)
    first = pages[0].get("text", "")[:500]
    middle = pages[len(pages) // 2].get("text", "")[:300]
    last = pages[-1].get("text", "")[-500:]
    return (len(pages), hash(first), hash(middle), hash(last))


def _printed_page_candidates(pages, printed_page):
    if printed_page is None:
        return []
    return [i for i, p in enumerate(pages) if not _toc_like_page(p["text"]) and _contains_printed_page_number(p["text"], printed_page)]


def _detect_section_start(pages, aliases):
    """Return (physical_index, printed_page), using TOC first and heading scan as fallback."""
    aliases = _heading_variants(aliases)
    toc_page = _toc_printed_page(pages, aliases)

    # Fast path: a Contents/Index entry gives the printed page. Among physical pages
    # carrying that printed number, pick the one that actually contains the heading.
    if toc_page is not None:
        candidates = _printed_page_candidates(pages, toc_page)
        best = None
        for i in candidates:
            score = max((_heading_score(pages[i]["text"], a, max_lines=320) for a in aliases), default=0)
            candidate = (score, -i, i)
            if best is None or candidate > best:
                best = candidate
        if best is not None and best[0] > 0:
            return best[2], toc_page
        if candidates:
            # Printed-page mapping is still stronger than a random body mention.
            return candidates[0], toc_page

    # Fallback for reports with no useful TOC: require an actual heading-level match.
    best = None
    for i, p in enumerate(pages):
        if _toc_like_page(p["text"]):
            continue
        score = max((_heading_score(p["text"], a, max_lines=320) for a in aliases), default=0)
        if score <= 0:
            continue
        candidate = (score, -i, i)
        if best is None or candidate > best:
            best = candidate
    return (best[2], None) if best else (None, None)


def build_section_map(pages):
    """Discover high-level section starts once and sort them in document order."""
    key = _doc_key(pages)
    if key in _SECTION_MAP_CACHE:
        return _SECTION_MAP_CACHE[key]

    found = []
    for label, aliases in BOUNDARY_SECTIONS:
        idx, printed = _detect_section_start(pages, aliases)
        if idx is None:
            continue
        found.append({
            "label": label,
            "aliases": aliases,
            "index": idx,
            "pdf_page": pages[idx].get("page"),
            "printed_page": printed,
        })

    # If multiple boundary labels land on one page, one physical boundary is enough.
    by_index = {}
    for item in found:
        idx = item["index"]
        old = by_index.get(idx)
        if old is None or len(_compact(item["label"])) > len(_compact(old["label"])):
            by_index[idx] = item

    result = sorted(by_index.values(), key=lambda x: x["index"])
    if len(_SECTION_MAP_CACHE) >= 24:
        _SECTION_MAP_CACHE.clear()
    _SECTION_MAP_CACHE[key] = result
    return result


def _next_boundary(pages, start_idx, printed_start=None):
    """End immediately before the next discovered document-level section."""
    for item in build_section_map(pages):
        if item["index"] > start_idx:
            printed_end = None
            if printed_start is not None and item.get("printed_page") is not None:
                printed_end = item["printed_page"] - 1
            return item["index"] - 1, printed_end
    return len(pages) - 1, None

def _preset_end(pages, start, printed_start, end_markers):
    """Return (physical_end_index, printed_end_page_or_None)."""

    # Best case: use the next section's page from the Contents/Index.
    toc_candidates = []
    if printed_start is not None:
        for marker in end_markers:
            p = _toc_printed_page(pages, marker)
            if p is None or p <= printed_start:
                continue
            idx = _physical_index_for_printed_page(pages, p, start_after=start)
            if idx is not None and idx > start:
                toc_candidates.append((idx, p))

        if toc_candidates:
            boundary_idx, boundary_printed = min(toc_candidates, key=lambda x: x[0])
            return boundary_idx - 1, boundary_printed - 1

    # Fallback: look only for heading-level boundary markers on later pages.
    for i in range(start + 1, len(pages)):
        if any(_heading_score(pages[i]["text"], marker, max_lines=220) > 0 for marker in end_markers):
            return i - 1, None

    return len(pages) - 1, None


def _payload(raw_pages, clean_pages_, start, end, printed_start=None, printed_end=None):
    raw_text = "\n\n".join(p["text"].strip() for p in raw_pages[start:end + 1]).strip()
    clean_text = "\n\n".join(p["text"].strip() for p in clean_pages_[start:end + 1]).strip()
    payload = {
        "start_page": raw_pages[start].get("page"),
        "end_page": raw_pages[end].get("page"),
        "text": clean_text,
        "raw_text": raw_text,
    }
    if printed_start is not None:
        payload["printed_start_page"] = printed_start
    if printed_end is not None:
        payload["printed_end_page"] = printed_end
    return payload


def extract_preset(raw_pages, clean_pages_, label):
    cfg = PRESETS[label]
    start, printed_start = _detect_section_start(clean_pages_, cfg["aliases"])
    if start is None:
        return None

    # Leadership pages often sit in the front matter where Contents and report page 1
    # can share a physical page. Their small, explicit end-marker lists are safer.
    if label in {"Chairman Message", "CEO Message", "Managing Director Message"}:
        end, printed_end = _preset_end(clean_pages_, start, printed_start, cfg["end"])
    else:
        end, printed_end = _next_boundary(clean_pages_, start, printed_start)

    if end < start:
        end, printed_end = start, printed_start

    return _payload(
        raw_pages, clean_pages_, start, end,
        printed_start=printed_start,
        printed_end=printed_end,
    )


def extract_custom(raw_pages, clean_pages_, heading):
    """Arbitrary section: exact TOC/heading start, then stop at the next major section."""
    heading = _norm_line(heading)
    if not heading:
        return None

    start, printed_start = _detect_section_start(clean_pages_, [heading])
    if start is None:
        return None

    end, printed_end = _next_boundary(clean_pages_, start, printed_start)
    if end < start:
        end, printed_end = start, printed_start

    return _payload(
        raw_pages, clean_pages_, start, end,
        printed_start=printed_start,
        printed_end=printed_end,
    )


def combine_sections(raw_pages, clean_pages_, sections, labels, combined_label):
    """Combine only the exact detected message sections.

    Important: do NOT take one continuous min-to-max page range, because unrelated
    pages may sit between Chairman and MD/CEO messages.
    """
    present = [(label, sections[label]) for label in labels if label in sections]
    if not present:
        return None

    clean_parts = []
    raw_parts = []
    page_ranges = []
    printed_ranges = []

    for label, sec in present:
        clean_parts.append(sec.get("text", "").strip())
        raw_parts.append(sec.get("raw_text", sec.get("text", "")).strip())

        if sec.get("start_page") is not None:
            page_ranges.append({
                "label": label,
                "start_page": sec.get("start_page"),
                "end_page": sec.get("end_page"),
            })

        if sec.get("printed_start_page") is not None:
            printed_ranges.append({
                "label": label,
                "start_page": sec.get("printed_start_page"),
                "end_page": sec.get("printed_end_page", sec.get("printed_start_page")),
            })

    starts = [r["start_page"] for r in page_ranges]
    ends = [r["end_page"] for r in page_ranges]

    payload = {
        "start_page": min(starts) if starts else None,
        "end_page": max(ends) if ends else None,
        "page_ranges": page_ranges,
        "text": "\n\n".join(x for x in clean_parts if x).strip(),
        "raw_text": "\n\n".join(x for x in raw_parts if x).strip(),
        "combined_from": [label for label, _ in present],
    }

    if printed_ranges:
        payload["printed_page_ranges"] = printed_ranges
        payload["printed_start_page"] = min(r["start_page"] for r in printed_ranges)
        payload["printed_end_page"] = max(r["end_page"] for r in printed_ranges)

    return payload


def raw_full_text(pages):
    if len(pages) == 1 and pages[0].get("page") is None:
        return pages[0]["text"].strip()
    return "\n\n".join(f"===== PAGE {p['page']} =====\n\n{p['text'].strip()}" for p in pages).strip()


def search_pages(pages, query, limit=50):
    query = query.strip()
    if not query:
        return []
    pat = re.compile(re.escape(query), re.I)
    hits = []
    for p in pages:
        for m in pat.finditer(p["text"]):
            a = max(0, m.start() - 90)
            b = min(len(p["text"]), m.end() + 150)
            hits.append({
                "page": p.get("page"),
                "snippet": re.sub(r"\s+", " ", p["text"][a:b]).strip(),
            })
            if len(hits) >= limit:
                return hits
    return hits


def base_stem(filename):
    stem = filename.rsplit(".", 1)[0]
    stem = re.sub(r"(?:_Annual_Report)?_Full_Text$", "", stem, flags=re.I)
    return stem


def infer_metadata(filename, pages):
    stem = base_stem(filename)

    # ---- Financial year detection ----
    # Prefer explicit annual-report labels near the front of the document.
    front = "\n".join(p.get("text", "")[:8000] for p in pages[:8])
    year = None

    explicit_patterns = [
        r"(?:integrated\s+)?annual\s+report(?:\s*&\s*accounts)?(?:\s+for\s+the\s+year)?\s*[:\-]?\s*(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})",
        r"annual\s+report.{0,80}?(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})",
        r"(?:financial\s+year|FY)\s*[:\-]?\s*(20\d{2})\s*[-–—/]\s*(20\d{2}|\d{2})",
    ]

    for pattern in explicit_patterns:
        m = re.search(pattern, front, re.I | re.S)
        if m:
            y1, y2 = m.group(1), m.group(2)
            year = f"{y1}-{y2[-2:]}"
            break

    # Older reports often say only "year ended 31st March, 2015".
    if not year:
        m = re.search(
            r"year\s+ended\s+(?:on\s+)?31(?:st)?\s+March[,\s]+(20\d{2})",
            front,
            re.I,
        )
        if m:
            end_year = int(m.group(1))
            year = f"{end_year - 1}-{str(end_year)[-2:]}"

    # Filename is a last resort only.
    if not year:
        m = re.search(r"\b(20\d{2})\s*[-–—_/]\s*(\d{2,4})\b", filename)
        if m:
            year = f"{m.group(1)}-{m.group(2)[-2:]}"
        else:
            m = re.search(r"\b(20\d{2})\b", filename)
            if m:
                year = m.group(1)

    # ---- Company name detection ----
    company = re.sub(r"\s*\(\d+\)\s*$", "", stem)
    company = re.sub(r"[_-]+", " ", company)
    company = re.sub(r"\b(?:Integrated\s+)?Annual\s+Report\b", " ", company, flags=re.I)
    company = re.sub(
        r"\b(?:Full\s*Text|Outputs?|Output|BRSR|BRR|ESG|MDA|"
        r"Management\s+Discussion(?:\s*&\s*Analysis)?|Chairman(?:'s)?\s+Message|"
        r"CEO\s+Message|MD\s+Message)\b",
        " ", company, flags=re.I,
    )
    company = re.sub(r"\b20\d{2}(?:\s*[-–—_/]?\s*\d{2,4})?\b", " ", company)
    company = re.sub(r"\s+", " ", company).strip(" _-") or stem.replace("_", " ").strip()

    # Generic filenames such as 2015.pdf / 2020.pdf: infer company from document text.
    if re.fullmatch(r"(?:20\d{2}|annual report|report|document)", company.strip(), re.I):
        text_sample = "\n".join(p.get("text", "") for p in pages[:20])
        candidates = re.findall(
            r"(?im)^\s*([A-Z][A-Za-z0-9&.,'()\- ]{2,100}\s(?:Limited|Ltd\.?))\s*$",
            text_sample
        )
        blocked = {
            "BSE Limited",
            "National Stock Exchange of India Limited",
            "Central Depository Services India Limited",
            "National Securities Depository Limited",
        }
        cleaned = []
        for c in candidates:
            c = re.sub(r"\s+", " ", c).strip()
            if c not in blocked:
                cleaned.append(c)
        if cleaned:
            company = Counter(cleaned).most_common(1)[0][0]
            company = re.sub(r"\s+Limited$", "", company, flags=re.I).strip()

    return {"company": company, "year": year or "Year not detected", "source_file": filename}



def _docx_bytes(title, text):
    d = Document()
    d.add_heading(title, 0)
    for block in text.split("\n\n"):
        if block.strip():
            d.add_paragraph(block.strip())
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _csv_bytes(rows):
    s = io.StringIO()
    w = csv.DictWriter(s, fieldnames=rows[0].keys())
    w.writeheader()
    w.writerows(rows)
    return s.getvalue().encode("utf-8-sig")


def _pdf_bytes(title, text):
    """Create a searchable text PDF using PyMuPDF Story (supports multi-page Unicode text)."""
    buf = io.BytesIO()
    writer = fitz.DocumentWriter(buf)
    mediabox = fitz.paper_rect("a4")
    rect = fitz.Rect(42, 42, mediabox.width - 42, mediabox.height - 42)
    body = html.escape(text).replace("\n", "<br>")
    story = fitz.Story(
        html=f"<h2>{html.escape(title)}</h2><div>{body}</div>",
        user_css="body {font-family: sans-serif; font-size: 9pt; line-height: 1.25;} h2 {font-size: 14pt;}"
    )
    more = True
    while more:
        dev = writer.begin_page(mediabox)
        more, _ = story.place(rect)
        story.draw(dev)
        writer.end_page()
    writer.close()
    return buf.getvalue()


def safe_label(label):
    return {
        "Full Report": "Annual_Report_Full_Text",
        "Leadership Messages (Combined)": "Chairman_CEO_MD_Messages",
        "Chairman Message": "Chairman_Message",
        "CEO Message": "CEO_Message",
        "Managing Director Message": "Managing_Director_Message",
        "Management Discussion & Analysis": "MDA",
        "Business Responsibility & Sustainability Report (BRSR)": "BRSR",
        "Business Responsibility Report (BRR)": "BRR",
        "ESG Report": "ESG_Report",
        "Sustainability Report": "Sustainability_Report",
    }.get(label, re.sub(r"[^A-Za-z0-9_-]+", "_", label).strip("_"))

# Backward-compatible internal alias.
_safe_label = safe_label


def format_file(stem, label, payload, fmt, raw_pages, source_bytes=None, source_is_pdf=False, all_sections=None):
    safe = safe_label(label)
    name = f"{stem}_{safe}"
    text = payload["text"]

    if fmt == "TXT":
        return name + ".txt", text.encode("utf-8")
    if fmt == "MD":
        return name + ".md", (f"# {label}\n\n{text}\n").encode("utf-8")
    if fmt == "DOCX":
        return name + ".docx", _docx_bytes(label, text)
    if fmt == "PDF":
        if label == "Full Report" and source_is_pdf and source_bytes is not None:
            return name + ".pdf", source_bytes
        return name + ".pdf", _pdf_bytes(label, text)
    if fmt == "JSON":
        if label == "Full Report":
            obj = {
                "file": stem,
                "page_count": len(raw_pages) if raw_pages and raw_pages[0].get("page") is not None else None,
                "pages": raw_pages,
                "text": text,
                "sections": {
                    k: {"start_page": v.get("start_page"), "end_page": v.get("end_page")}
                    for k, v in (all_sections or {}).items()
                },
            }
        else:
            obj = payload
        return name + ".json", json.dumps(obj, ensure_ascii=False, indent=2).encode("utf-8")
    if fmt == "CSV":
        if label == "Full Report":
            rows = [{"page": p.get("page"), "text": p["text"]} for p in raw_pages]
        else:
            rows = [{
                "section": label,
                "start_page": payload.get("start_page"),
                "end_page": payload.get("end_page"),
                "text": text,
            }]
        return name + ".csv", _csv_bytes(rows)
    raise ValueError(fmt)


def make_zip(files):
    b = io.BytesIO()
    with zipfile.ZipFile(b, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    return b.getvalue()


def manifest_csv(results):
    s = io.StringIO()
    fields = ["company", "year", "source_file", "section", "start_page", "end_page", "formats"]
    w = csv.DictWriter(s, fieldnames=fields)
    w.writeheader()
    for r in results:
        w.writerow({
            "company": r["meta"]["company"], "year": r["meta"]["year"],
            "source_file": r["name"], "section": "Full Report",
            "start_page": 1 if len(r["raw_pages"]) > 1 else "",
            "end_page": len(r["raw_pages"]) if len(r["raw_pages"]) > 1 else "",
            "formats": ", ".join(r["formats"]),
        })
        for label, sec in r["sections"].items():
            w.writerow({
                "company": r["meta"]["company"], "year": r["meta"]["year"],
                "source_file": r["name"], "section": label,
                "start_page": sec.get("start_page", ""), "end_page": sec.get("end_page", ""),
                "formats": ", ".join(r["formats"]),
            })
    return s.getvalue().encode("utf-8-sig")


def clean_background_pdf(pdf_bytes, threshold=215):
    """Create a separate grayscale background-clean copy; original is never modified."""
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    out = fitz.open()
    lut = [x if x < threshold else 255 for x in range(256)]
    for p in src:
        pix = p.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), colorspace=fitz.csGRAY, alpha=False)
        img = Image.frombytes("L", (pix.width, pix.height), pix.samples).point(lut)
        b = io.BytesIO()
        img.save(b, "PNG")
        q = out.new_page(width=p.rect.width, height=p.rect.height)
        q.insert_image(q.rect, stream=b.getvalue())
    return out.tobytes(deflate=True)
